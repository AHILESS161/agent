"""Формирование чернового заявления на регистрацию товарного знака.

Заполняется официальный бланк (приложение № 1 к Требованиям) в формате
DOCX. Исходный формат выбран намеренно: генерировать документ и затем
сохранять его в PDF надёжнее, чем разбирать заполненный PDF обратно.

Главное правило: в документ попадают ТОЛЬКО подтверждённые специалистом
значения. Поле со статусом missing, conflict или needs_review остаётся
пустым, а причина указывается в чек-листе. Черновик юридически значимого
документа не должен содержать непроверенных данных, даже если система
в них уверена.
"""

from __future__ import annotations

import hashlib
import io
import re
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import docx
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from sqlalchemy import func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.logging import get_logger
from app.infrastructure.database.models import (
    ApplicationDraft,
    DraftStatus,
    ExtractedField,
    FieldStatus,
    SourceDocument,
    TrademarkApplicationDraft,
)
from app.document_processing.mappers.field_mapping import FieldMappingEngine
from app.services.class_analysis import load_class_context
from app.services.file_storage import read_file

logger = get_logger(__name__)

# Вид знака в терминах бланка: коды 550–558 и 551.
_MARK_KIND_LABELS = {
    "word": "словесный",
    "figurative": "изобразительный",
    "combined": "комбинированный",
    "3d": "объёмный (554)",
    "sound": "звуковой (556)",
    "color": "цветовой (558)",
    "other": "иной",
}


def _has_foreign_wording(value: str | None) -> bool:
    """Нужны ли в заявлении перевод/транслитерация иностранного элемента."""
    return bool(re.search(r"[A-Za-z]", value or ""))

TEMPLATE_PATH = (
    Path(__file__).resolve().parents[2]
    / "resources"
    / "application_templates"
    / "trademark_application_blank.docx"
)

SCHEMA_VERSION = "1.0.0"


@dataclass
class FilledField:
    field_id: str
    label: str
    value: str
    source: str


@dataclass
class SkippedField:
    field_id: str
    label: str
    reason: str


@dataclass
class DraftContent:
    """Данные для заполнения бланка."""

    filled: list[FilledField] = field(default_factory=list)
    skipped: list[SkippedField] = field(default_factory=list)
    checklist: list[str] = field(default_factory=list)
    # Перечень товаров для кода 511: (номер класса, наименование).
    # Только подтверждённые специалистом классы: неподтверждённый
    # перечень в заявление попадать не должен.
    classes: list[tuple[str, str]] = field(default_factory=list)

    def value_of(self, field_id: str) -> str | None:
        for item in self.filled:
            if item.field_id == field_id:
                return item.value
        return None


async def collect_draft_content(
    session: AsyncSession, application: TrademarkApplicationDraft
) -> DraftContent:
    """Собрать значения для заявления из подтверждённых полей дела."""
    content = DraftContent()

    rows = (
        (
            await session.execute(
                select(ExtractedField).where(
                    ExtractedField.application_id == application.id
                )
            )
        )
        .scalars()
        .all()
    )
    by_path = {row.field_path: row for row in rows}

    client = application.client
    case_values = {
        "case.applicant.full_name": client.full_name_or_company_name if client else None,
        "case.applicant.inn": client.inn if client else None,
        "case.applicant.ogrn": client.ogrn_or_ogrnip if client else None,
        "case.applicant.kpp": client.kpp if client else None,
        "case.applicant.legal_address": client.address if client else None,
        "case.applicant.country_code": (client.country or "RU") if client else "RU",
    }

    engine = FieldMappingEngine()
    for spec in engine.config.get("mappings", []):
        target = spec.get("application_field")
        if not target:
            continue

        label = spec.get("label", target)
        registry_field = spec.get("registry_field")
        required = bool(spec.get("required_for_application", False))

        row = by_path.get(registry_field) if registry_field else None
        case_value = case_values.get(spec.get("case_field"))

        if row is None:
            if case_value:
                # Пользователь уже проверил и сохранил значение в карточке
                # заявителя. Оно не обязано присутствовать в выписке (адрес
                # ИП, например, обычно скрыт в открытой ЕГРИП).
                content.filled.append(
                    FilledField(
                        field_id=target,
                        label=label,
                        value=str(case_value),
                        source="введено пользователем",
                    )
                )
            elif spec.get("default_value"):
                # Значение по умолчанию — предложение, а не факт.
                content.skipped.append(
                    SkippedField(
                        field_id=target,
                        label=label,
                        reason=(
                            f"Значение по умолчанию «{spec['default_value']}» "
                            "не подтверждено специалистом"
                        ),
                    )
                )
            elif required:
                content.skipped.append(
                    SkippedField(
                        field_id=target,
                        label=label,
                        reason="Значение не извлечено из документов",
                    )
                )
            continue

        if row.status is FieldStatus.confirmed:
            content.filled.append(
                FilledField(
                    field_id=target,
                    label=label,
                    value=row.normalized_value or row.raw_value or "",
                    source=(
                        f"{row.extraction_method.value}"
                        + (f", стр. {row.page_number}" if row.page_number else "")
                    ),
                )
            )
            continue

        # Всё, что не подтверждено, в документ не попадает.
        reasons = {
            FieldStatus.conflict: "несколько несовпадающих значений, выбор не сделан",
            FieldStatus.needs_review: "значение требует проверки специалистом",
            FieldStatus.missing: "значение не найдено в документах",
            FieldStatus.matched: "значение извлечено, но не подтверждено специалистом",
            FieldStatus.rejected: "значение отклонено специалистом",
            FieldStatus.left_empty: "специалист оставил поле пустым",
        }
        content.skipped.append(
            SkippedField(
                field_id=target,
                label=label,
                reason=reasons.get(row.status, row.status.value),
            )
        )

    # --- сведения об обозначении берутся из карточки дела ---
    mark_text = application.mark_text or application.mark_name
    if mark_text:
        content.filled.append(
            FilledField(
                field_id="application.mark.text",
                label="Заявляемое обозначение",
                value=mark_text,
                source="карточка дела",
            )
        )
    else:
        content.skipped.append(
            SkippedField(
                field_id="application.mark.text",
                label="Заявляемое обозначение",
                reason="не указано в деле",
            )
        )

    # Вид знака отмечается в бланке галочкой, но специалисту нужно
    # видеть выбранное значение и в перечне полей.
    if application.mark_type:
        content.filled.append(
            FilledField(
                field_id="application.mark.kind",
                label="Вид знака",
                value=_MARK_KIND_LABELS.get(
                    application.mark_type.value, application.mark_type.value
                ),
                source="карточка дела",
            )
        )

    if application.description_of_mark:
        content.filled.append(
            FilledField(
                field_id="application.mark.description",
                label="Описание обозначения",
                value=application.description_of_mark,
                source="карточка дела",
            )
        )

    # Контакты и адрес для переписки берутся из данных заявителя. Для
    # самостоятельной подачи это ожидаемый вариант; отдельный адрес можно
    # будет добавить позднее, если он отличается.
    if client:
        for field_id, label, value in (
            ("application.correspondence_address", "Адрес для переписки", client.address),
            ("application.contact.phone", "Телефон для переписки", client.phone),
            ("application.contact.email", "E-mail для переписки", client.email),
        ):
            if value:
                content.filled.append(
                    FilledField(
                        field_id=field_id,
                        label=label,
                        value=value,
                        source="данные заявителя",
                    )
                )

    language_fields = (
        (
            "application.mark.transliteration",
            "Транслитерация",
            application.transliteration,
        ),
        ("application.mark.translation", "Перевод", application.translation),
    ) if _has_foreign_wording(mark_text) else ()

    for field_id, label, value in (
        ("application.mark.colors", "Цвет или цветовое сочетание", application.colors_claimed),
        *language_fields,
        ("application.signatory.name", "ФИО подписанта", application.signatory_name),
        ("application.signatory.position", "Должность подписанта", application.signatory_position),
        (
            "application.signatory.date",
            "Дата подписания",
            application.signature_date.strftime("%d.%m.%Y")
            if application.signature_date
            else None,
        ),
    ):
        if value:
            content.filled.append(
                FilledField(
                    field_id=field_id,
                    label=label,
                    value=value,
                    source="данные заявки",
                )
            )

    # --- классы МКТУ ---
    class_context = await load_class_context(session, application.id)
    if class_context.is_confirmed:
        for item in class_context.approved:
            content.filled.append(
                FilledField(
                    field_id=f"application.goods_services.class_{item.class_number}",
                    label=f"Класс {item.class_number}",
                    value=item.class_description or "",
                    source="подтверждено специалистом",
                )
            )
            content.classes.append(
                (str(item.class_number), item.class_description or "")
            )
    elif class_context.has_any:
        content.skipped.append(
            SkippedField(
                field_id="application.goods_services",
                label="Перечень товаров и услуг",
                reason="классы МКТУ не подтверждены специалистом",
            )
        )
    else:
        content.skipped.append(
            SkippedField(
                field_id="application.goods_services",
                label="Перечень товаров и услуг",
                reason="классы МКТУ не определены",
            )
        )

    # --- чек-лист недостающего ---
    content.checklist = _build_checklist(content, application)
    return content


def _build_checklist(
    content: DraftContent, application: TrademarkApplicationDraft
) -> list[str]:
    """Что специалисту нужно сделать перед подачей."""
    checklist: list[str] = []

    for item in content.skipped:
        checklist.append(f"{item.label}: {item.reason}")

    # Поля, которые заведомо не берутся из документов.
    if not application.mark_image_file_id:
        checklist.append(
            "Изображение обозначения (540): файл не приложен к делу"
        )
    if not application.signatory_name:
        checklist.append("Укажите ФИО человека, который подпишет заявление")
    if not application.signature_date:
        checklist.append("Укажите дату подписания заявления")
    checklist.append(
        "При бумажной подаче поставьте собственноручную подпись в готовом бланке; "
        "при электронной подаче подпишите отправление электронной подписью в "
        "официальном сервисе Роспатента"
    )
    checklist.append(
        "После оплаты проверьте реквизиты платежа в официальном сервисе; "
        "подтверждение оплаты прикладывается по инициативе заявителя"
    )
    return checklist


# ---------------------------------------------------------------------------
# Заполнение бланка
# ---------------------------------------------------------------------------

# Куда писать значения: якорь в бланке -> идентификатор поля.
# Ячейка ищется по началу текста, значение дописывается в неё же.
_ANCHOR_MAP: list[tuple[str, str]] = [
    ("(731) ЗАЯВИТЕЛЬ", "application.applicant.name"),
    ("ИДЕНТИФИКАТОРЫ ЗАЯВИТЕЛЯ", "application.applicant.identifiers"),
    ("(540) ЗАЯВЛЯЕМОЕ ОБОЗНАЧЕНИЕ", "application.mark.text"),
]


def _template_hash() -> str | None:
    if not TEMPLATE_PATH.exists():
        return None
    return hashlib.sha256(TEMPLATE_PATH.read_bytes()).hexdigest()


def _find_cell(table, marker: str):
    """Найти ячейку бланка по началу её текста.

    Бланк — одна таблица 118×44 с объединёнными ячейками, поэтому
    позиции ищутся по подписи поля, а не по фиксированным индексам:
    так разметка переживёт мелкие правки формы.
    """
    normalized = marker.replace(" ", " ")
    for row in table.rows:
        for cell in row.cells:
            text = cell.text.strip().replace(" ", " ")
            if text.startswith(normalized):
                return cell
    return None


def _write_into(cell, lines: list[str]) -> bool:
    """Дописать значение в ячейку бланка под её подписью.

    В бланке под каждой подписью оставлен пустой абзац — туда и
    пишется значение. Если пустого абзаца нет, добавляется новый:
    затирать типографский текст формы нельзя.
    """
    if cell is None or not lines:
        return False

    target = None
    for paragraph in cell.paragraphs[1:]:
        if not paragraph.text.strip():
            target = paragraph
            break
    if target is None:
        target = cell.add_paragraph()

    target.add_run("\n".join(lines)).bold = True
    return True


def _check_box(table, row_index: int, cell_index: int) -> bool:
    """Поставить однозначную отметку X в квадрате официального бланка."""
    if row_index >= len(table.rows) or cell_index >= len(table.rows[row_index].cells):
        return False
    cell = table.rows[row_index].cells[cell_index]
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("X")
    run.bold = True
    run.font.size = Pt(9)
    return True


def _fill_goods_table(table, classes: list[tuple[str, str]]) -> bool:
    """Заполнить перечень товаров: класс и наименование (код 511)."""
    header = None
    for index, row in enumerate(table.rows):
        if row.cells[0].text.strip() == "Класс":
            header = index
            break
    if header is None:
        return False

    written = 0
    for offset, (number, goods) in enumerate(classes, start=1):
        position = header + offset
        if position >= len(table.rows):
            break
        row = table.rows[position]
        _write_into(row.cells[0], [number]) or row.cells[0].paragraphs[0].add_run(
            number
        )
        # Наименование товаров занимает правую часть строки.
        value_cell = row.cells[3] if len(row.cells) > 3 else None
        if value_cell is not None:
            _write_into(value_cell, [goods]) or value_cell.paragraphs[0].add_run(goods)
        written += 1
    return written > 0


def _fill_mark_block(
    cell,
    *,
    mark_text: str,
    description_lines: list[str],
) -> None:
    """Заполнить объединённый блок (540)/(571) двумя реальными колонками.

    В официальном шаблоне левая и правая области обозначены табуляцией внутри
    одной объединённой ячейки. Добавление обычного абзаца помещает содержимое
    ниже обеих подписей. Вложенная таблица без служебного текста сохраняет
    требуемую геометрию: обозначение слева, описание справа.
    """
    block = cell.add_table(rows=1, cols=2)
    block.autofit = False
    left, right = block.rows[0].cells
    left.width = Inches(3.45)
    right.width = Inches(3.45)

    left_paragraph = left.paragraphs[0]
    left_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if mark_text:
        run = left_paragraph.add_run(mark_text)
        run.font.size = Pt(30)
        run.bold = True

    right_paragraph = right.paragraphs[0]
    if description_lines:
        right_paragraph.add_run("\n".join(description_lines))


def _fill_mark_image_box(
    table,
    mark_image: bytes,
    *,
    application_id: int | None,
) -> None:
    """Вставить изображение в специальный квадрат под полем (540)."""
    # В официальном шаблоне это строка с фиксированной высотой 3856 twips,
    # а квадрат занимает колонки 1–15. Он не имеет текстового якоря, поэтому
    # адресуется по устойчивой структуре самой утверждённой формы.
    if len(table.rows) <= 19 or len(table.rows[19].cells) <= 1:
        return
    box = table.rows[19].cells[1]
    paragraph = box.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    box.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    try:
        from PIL import Image, ImageChops

        with Image.open(io.BytesIO(mark_image)) as source:
            image = source.convert("RGBA")
            # У логотипов часто остаются большие белые/прозрачные поля. Если
            # вставлять исходный холст, сам знак выглядит маленьким и визуально
            # смещённым внутри предназначенного для него квадрата формы.
            alpha = image.getchannel("A")
            alpha_bbox = alpha.getbbox()
            rgb = image.convert("RGB")
            difference = ImageChops.difference(
                rgb, Image.new("RGB", rgb.size, (255, 255, 255))
            ).convert("L")
            difference = difference.point(lambda value: 255 if value > 12 else 0)
            content_bbox = difference.getbbox() or alpha_bbox
            if content_bbox:
                left, top, right, bottom = content_bbox
                padding = max(6, int(min(image.size) * 0.015))
                image = image.crop(
                    (
                        max(0, left - padding),
                        max(0, top - padding),
                        min(image.width, right + padding),
                        min(image.height, bottom + padding),
                    )
                )
            ratio = image.width / max(image.height, 1)
            normalized = io.BytesIO()
            image.save(normalized, format="PNG")
        max_width, max_height = 2.65, 2.42
        width = min(max_width, max_height * ratio)
        height = width / max(ratio, 0.01)
        paragraph.add_run().add_picture(
            io.BytesIO(normalized.getvalue()),
            width=Inches(width),
            height=Inches(height),
        )
    except Exception as exc:  # noqa: BLE001
        logger.warning(
            "Не удалось вставить изображение обозначения в квадрат DOCX",
            application_id=application_id,
            error=str(exc),
        )


def _fill_attachment_row(
    table,
    row_index: int,
    *,
    sheets: int = 1,
    copies: int = 1,
) -> None:
    """Отметить приложение и заполнить количество листов/экземпляров."""
    _check_box(table, row_index, 1)
    for cell_index, value in ((37, sheets), (43, copies)):
        if row_index < len(table.rows) and cell_index < len(table.rows[row_index].cells):
            cell = table.rows[row_index].cells[cell_index]
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.add_run(str(value)).bold = True


def _fill_signature_block(
    table,
    *,
    name: str | None,
    position: str | None,
    signed_at: str | None,
    filing_method: str,
) -> None:
    """Заполнить сведения о подписанте, не имитируя собственноручную подпись.

    Для бумажной подачи в поле остаётся линия для подписи от руки. Для
    электронной подачи юридически значимая подпись ставится в официальном
    сервисе, поэтому изображение росчерка в DOCX не создаётся.
    """
    if len(table.rows) <= 117:
        return
    details = ", ".join(part for part in (name, position) if part)
    if details:
        paragraph = table.rows[116].cells[0].paragraphs[0]
        if filing_method == "paper":
            paragraph.add_run("________________ / ")
        paragraph.add_run(details).bold = True
    if signed_at:
        cell = table.rows[117].cells[0]
        paragraph = cell.paragraphs[-1]
        paragraph.add_run(signed_at).bold = True


def _safe_claimed_colors(claimed: str | None, mark_image: bytes | None) -> str | None:
    """Не переносить в заявление цвет, противоречащий самому изображению."""
    if not claimed or not mark_image:
        return claimed
    try:
        from PIL import Image, ImageStat

        with Image.open(io.BytesIO(mark_image)).convert("RGB") as image:
            image.thumbnail((256, 256))
            saturation = ImageStat.Stat(
                image.convert("HSV").getchannel("S")
            ).mean[0]
        if saturation < 12:
            normalized = claimed.casefold()
            neutral = ("черн", "бел", "сер", "black", "white", "gray", "grey")
            if not any(token in normalized for token in neutral):
                logger.warning(
                    "Заявленный цвет не перенесён: изображение фактически монохромное",
                    claimed=claimed,
                )
                return None
    except Exception as exc:  # noqa: BLE001
        logger.warning("Не удалось проверить соответствие заявленного цвета", error=str(exc))
    return claimed


async def load_mark_image_content(
    session: AsyncSession, application: TrademarkApplicationDraft
) -> bytes | None:
    """Загрузить активное изображение обозначения из защищённого хранилища."""
    raw_id = application.mark_image_file_id
    if not raw_id or not str(raw_id).isdigit():
        return None
    document = (
        await session.execute(
            select(SourceDocument).where(
                SourceDocument.id == int(raw_id),
                SourceDocument.application_id == application.id,
            )
        )
    ).scalar_one_or_none()
    if document is None:
        return None
    try:
        return read_file(document.stored_path)
    except (FileNotFoundError, ValueError):
        logger.warning(
            "Активное изображение обозначения не найдено в хранилище",
            application_id=application.id,
            document_id=document.id,
        )
        return None


def render_docx(
    content: DraftContent,
    application: TrademarkApplicationDraft,
    mark_image: bytes | None = None,
    include_goods_attachment: bool = False,
) -> bytes:
    """Заполнить официальный бланк заявки подтверждёнными данными.

    Заполняется именно бланк Роспатента (приложение к приказу
    Минэкономразвития), а не собственная форма: документ должен быть
    пригоден к подаче, а не пересказывать её своими словами.

    В документ не добавляется никаких служебных пометок — ни о том,
    что это черновик, ни об использовании AI. Это бланк заявления,
    и посторонний текст в нём недопустим. Предупреждения, перечень
    незаполненных полей и чек-лист живут в интерфейсе, а выгрузка
    закрыта до утверждения специалистом.

    Незаполненные поля остаются пустыми — ровно как в бумажной форме,
    которую специалист дозаполняет вручную.
    """
    if not TEMPLATE_PATH.exists():
        raise FileNotFoundError(f"Бланк заявления не найден: {TEMPLATE_PATH}")

    document = docx.Document(str(TEMPLATE_PATH))

    # Служебная надпись «Приложение № 1…» относится к публикации формы,
    # а не к заполняемому заявлению. В выгружаемый пользователем документ она
    # и содержимое колонтитулов не попадают.
    for paragraph in list(document.paragraphs):
        if paragraph.text.strip().startswith("Приложение № 1"):
            paragraph._element.getparent().remove(paragraph._element)
    for section in document.sections:
        for container in (section.header, section.footer):
            for paragraph in container.paragraphs:
                paragraph.clear()

    table = document.tables[0]

    values = {item.field_id: item.value for item in content.filled}

    # --- (731) заявитель: наименование и адрес ---
    applicant_block = [
        values.get("application.applicant.name", ""),
        values.get("application.applicant.address", ""),
    ]
    _write_into(
        _find_cell(table, "(731)"), [line for line in applicant_block if line]
    )

    # --- идентификаторы заявителя ---
    identifiers: list[str] = []
    for label, field_id in (
        ("ОГРН", "application.applicant.ogrn"),
        ("ИНН", "application.applicant.inn"),
        ("КПП", "application.applicant.kpp"),
    ):
        value = values.get(field_id)
        if value:
            identifiers.append(f"{label}: {value}")
    country_code = values.get("application.applicant.country_code")
    if country_code:
        identifiers.append(f"Код страны (ВОИС ST.3): {country_code}")
    _write_into(_find_cell(table, "ИДЕНТИФИКАТОРЫ"), identifiers)

    # --- (750) адрес и контакты для переписки ---
    correspondence = [
        values.get("application.applicant.name", ""),
        values.get("application.correspondence_address", ""),
        f"Телефон: {values['application.contact.phone']}" if values.get("application.contact.phone") else "",
        f"E-mail: {values['application.contact.email']}" if values.get("application.contact.email") else "",
    ]
    _write_into(_find_cell(table, "(750)"), [line for line in correspondence if line])

    # --- (540) заявляемое обозначение ---
    mark_text = values.get("application.mark.text") or (
        application.mark_text or application.mark_name or ""
    )
    # В поле (540) должно быть одно заявляемое обозначение. Для
    # изобразительного и комбинированного знака это загруженная картинка в
    # предназначенном формой квадрате; отдельное повторение слов сверху
    # создаёт впечатление, что заявляются два обозначения.
    if mark_image and getattr(application.mark_type, "value", None) in {
        "figurative",
        "combined",
    }:
        mark_text_for_form = ""
    else:
        mark_text_for_form = mark_text
    include_language = _has_foreign_wording(mark_text)
    description_lines = [
        values.get("application.mark.description", ""),
        f"Транслитерация: {values['application.mark.transliteration']}" if include_language and values.get("application.mark.transliteration") else "",
        f"Перевод: {values['application.mark.translation']}" if include_language and values.get("application.mark.translation") else "",
    ]
    anchor = _find_cell(table, "(540)")
    if anchor is not None:
        _fill_mark_block(
            anchor,
            mark_text=mark_text_for_form,
            description_lines=[line for line in description_lines if line],
        )
    if mark_image:
        _fill_mark_image_box(
            table,
            mark_image,
            application_id=getattr(application, "id", None),
        )

    colors = _safe_claimed_colors(values.get("application.mark.colors"), mark_image)
    if colors:
        _write_into(_find_cell(table, "(591)"), [colors])
        _check_box(table, 23, 1)

    # Вид знака известен системе и отмечается в самом квадрате формы, а не
    # повторяется свободным текстом под заголовком (550).
    mark_type = (
        application.mark_type.value
        if getattr(application, "mark_type", None) is not None
        else None
    )
    mark_type_boxes = {
        "word": (26, 15),
        "figurative": (26, 26),
        "3d": (30, 1),
        "sound": (30, 23),
        "color": (32, 1),
        "combined": (34, 1),
        "other": (36, 1),
    }
    if mark_type in mark_type_boxes:
        _check_box(table, *mark_type_boxes[mark_type])

    # Загруженное обозначение входит в состав электронной заявки.
    if mark_image or mark_type in {"sound", "3d"}:
        _fill_attachment_row(table, 90)

    # Просьба о бумажном свидетельстве — отдельный осознанный выбор
    # заявителя. По умолчанию остаётся электронное свидетельство.
    if getattr(application, "request_paper_certificate", False):
        _check_box(table, 83, 1)

    # --- (511) перечень товаров и услуг ---
    if content.classes:
        _fill_goods_table(table, content.classes)
        if include_goods_attachment:
            _fill_attachment_row(table, 93)

    _fill_signature_block(
        table,
        name=(
            values.get("application.signatory.name")
            or getattr(application, "signatory_name", None)
        ),
        position=(
            values.get("application.signatory.position")
            or getattr(application, "signatory_position", None)
        ),
        signed_at=(
            values.get("application.signatory.date")
            or (
                application.signature_date.strftime("%d.%m.%Y")
                if getattr(application, "signature_date", None)
                else None
            )
        ),
        filing_method=getattr(application, "filing_method", None) or "electronic",
    )

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def create_draft(
    session: AsyncSession,
    application: TrademarkApplicationDraft,
    user_id: int | None = None,
) -> ApplicationDraft:
    """Сформировать новую версию чернового заявления."""
    content = await collect_draft_content(session, application)
    payload = render_docx(
        content,
        application,
        mark_image=await load_mark_image_content(session, application),
    )

    from app.services import file_storage

    stored = file_storage.save_upload(
        payload, f"application-draft-{application.id}.docx"
    )

    last_version = (
        await session.execute(
            select(func.max(ApplicationDraft.version)).where(
                ApplicationDraft.application_id == application.id
            )
        )
    ).scalar() or 0

    draft = ApplicationDraft(
        application_id=application.id,
        version=last_version + 1,
        status=DraftStatus.draft,
        filled_fields_json=[
            {
                "field_id": item.field_id,
                "label": item.label,
                "value": item.value,
                "source": item.source,
            }
            for item in content.filled
        ],
        skipped_fields_json=[
            {"field_id": item.field_id, "label": item.label, "reason": item.reason}
            for item in content.skipped
        ],
        checklist_json=content.checklist,
        file_path=stored.stored_path,
        file_sha256=stored.sha256,
        template_name=TEMPLATE_PATH.name,
        template_sha256=_template_hash(),
        schema_version=SCHEMA_VERSION,
        mapping_version=FieldMappingEngine().version,
        created_by_user_id=user_id,
    )
    session.add(draft)
    await session.flush()

    logger.info(
        "Черновик заявления сформирован",
        application_id=application.id,
        draft_id=draft.id,
        version=draft.version,
        filled=len(content.filled),
        skipped=len(content.skipped),
    )
    return draft


def serialize_draft(draft: ApplicationDraft) -> dict[str, Any]:
    return {
        "id": draft.id,
        "application_id": draft.application_id,
        "version": draft.version,
        "status": draft.status.value,
        "filled_fields": draft.filled_fields_json or [],
        "skipped_fields": draft.skipped_fields_json or [],
        "checklist": draft.checklist_json or [],
        "file_sha256": draft.file_sha256,
        "provenance": {
            "template_name": draft.template_name,
            "template_sha256": draft.template_sha256,
            "schema_version": draft.schema_version,
            "mapping_version": draft.mapping_version,
        },
        "approved_by_user_id": draft.approved_by_user_id,
        "approved_at": draft.approved_at.isoformat() if draft.approved_at else None,
        "exported_at": draft.exported_at.isoformat() if draft.exported_at else None,
        "created_at": draft.created_at.isoformat() if draft.created_at else None,
        # Экспорт разрешён только после утверждения специалистом.
        "can_export": draft.status
        in (DraftStatus.approved_by_specialist, DraftStatus.exported),
        "notice": (
            "Черновик содержит только подтверждённые значения. "
            "Поля со статусом «конфликт», «требует проверки» или «не найдено» "
            "оставлены пустыми намеренно."
        ),
    }
