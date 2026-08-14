"""
Утилиты извлечения текста из документов приложенной формы заявки.

Поддерживает:
    - PDF  (через pdfplumber)
    - DOCX (через python-docx)
    - TXT  (как fallback, на случай если pdfplumber не справился)
    - сканы PDF, PNG и JPG (OCR через Tesseract, русский + английский)

Возвращает чистый текст, пригодный для эвристического и LLM-парсинга.
"""

from __future__ import annotations

import io
import logging
import os
from dataclasses import dataclass
from pathlib import Path

from app.core.config import settings

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Lazy imports
# ---------------------------------------------------------------------------

try:
    import pdfplumber
    _PDFPLUMBER_AVAILABLE = True
except ImportError:
    _PDFPLUMBER_AVAILABLE = False

try:
    import docx as _docx  # python-docx
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image, ImageOps
    _OCR_IMPORTS_AVAILABLE = True
except ImportError:
    _OCR_IMPORTS_AVAILABLE = False


class UnsupportedDocumentType(ValueError):
    """Не удалось распознать формат документа."""


def detect_extension(filename: str) -> str:
    name = filename.lower()
    if name.endswith(".pdf"):
        return ".pdf"
    if name.endswith(".docx"):
        return ".docx"
    if name.endswith(".doc"):
        return ".doc"
    if name.endswith(".txt"):
        return ".txt"
    raise UnsupportedDocumentType(f"Unsupported file type: {filename}")


def extract_text_from_bytes(content: bytes, filename: str) -> str:
    """Главная точка входа: принимает байты и имя файла, возвращает текст."""
    ext = detect_extension_for_pages(filename)
    if ext == ".pdf":
        return extract_pdf_text(content)
    if ext in (".docx", ".doc"):
        return extract_docx_text(content)
    if ext == ".txt":
        return content.decode("utf-8", errors="replace")
    if ext in (".png", ".jpg", ".jpeg"):
        return "\n\n".join(
            page.text for page in extract_pages_from_bytes(content, filename)
            if page.text.strip()
        )
    raise UnsupportedDocumentType(f"Unsupported extension: {ext}")


def extract_pdf_text(content: bytes) -> str:
    """Извлекает текст из PDF; для скана автоматически применяет OCR."""
    pages = extract_pages_from_bytes(content, "document.pdf")
    result = "\n\n".join(page.text for page in pages if page.text.strip())
    logger.info("extract_pdf_text: %d pages, %d chars", len(pages), len(result))
    return result


def extract_docx_text(content: bytes) -> str:
    """Извлекает текст из DOCX через python-docx.

    Проходим по параграфам и таблицам — таблицы важны для бланков Роспатента,
    потому что часть полей там находится именно в табличной форме.
    """
    if not _DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    doc = _docx.Document(io.BytesIO(content))
    chunks: list[str] = []

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            chunks.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = (cell.text or "").strip()
                if cell_text:
                    chunks.append(cell_text)

    result = "\n".join(chunks)
    logger.info("extract_docx_text: %d chunks, %d chars", len(chunks), len(result))
    return result


@dataclass(frozen=True)
class ExtractedPage:
    """Текст одной страницы с указанием способа получения."""

    page_number: int
    text: str
    method: str            # значение ExtractionMethod
    ocr_confidence: float | None = None


def _prepare_image(image):
    """Нормализовать изображение, не увеличивая бесконтрольно память."""
    image = ImageOps.exif_transpose(image).convert("L")
    image = ImageOps.autocontrast(image)
    pixels = image.width * image.height
    if pixels > settings.OCR_MAX_IMAGE_PIXELS:
        scale = (settings.OCR_MAX_IMAGE_PIXELS / pixels) ** 0.5
        image = image.resize(
            (max(1, int(image.width * scale)), max(1, int(image.height * scale)))
        )
    return image


def _ocr_image(image) -> tuple[str, float | None]:
    """Распознать изображение и вернуть текст вместе со средней уверенностью."""
    if not settings.OCR_ENABLED:
        raise NoTextLayerError("OCR отключён в настройках сервера.")
    if not _OCR_IMPORTS_AVAILABLE:
        raise NoTextLayerError(
            "OCR недоступен: не установлен Python-пакет pytesseract."
        )

    if settings.OCR_TESSERACT_CMD:
        pytesseract.pytesseract.tesseract_cmd = settings.OCR_TESSERACT_CMD

    prepared = _prepare_image(image)
    tesseract_config = f"--oem 3 --psm {settings.OCR_PSM}"
    if settings.OCR_TESSDATA_DIR:
        os.environ["TESSDATA_PREFIX"] = settings.OCR_TESSDATA_DIR
    try:
        data = pytesseract.image_to_data(
            prepared,
            lang=settings.OCR_LANGUAGES,
            config=tesseract_config,
            output_type=pytesseract.Output.DICT,
            timeout=settings.OCR_TIMEOUT_SECONDS,
        )
    except pytesseract.TesseractNotFoundError as exc:
        raise NoTextLayerError(
            "OCR недоступен: Tesseract не установлен или не найден."
        ) from exc
    except RuntimeError as exc:
        raise NoTextLayerError(f"OCR не завершён: {exc}") from exc

    lines: list[str] = []
    current_key: tuple[int, int, int] | None = None
    current_tokens: list[str] = []
    weighted_confidence = 0.0
    confidence_weight = 0

    for index, raw_text in enumerate(data.get("text", [])):
        token = str(raw_text or "").strip()
        if not token:
            continue
        key = (
            int(data["block_num"][index]),
            int(data["par_num"][index]),
            int(data["line_num"][index]),
        )
        if current_key is not None and key != current_key and current_tokens:
            lines.append(" ".join(current_tokens))
            current_tokens = []
        current_key = key
        current_tokens.append(token)

        try:
            confidence = float(data["conf"][index])
        except (KeyError, TypeError, ValueError):
            confidence = -1
        if confidence >= 0:
            weight = max(1, len(token))
            weighted_confidence += confidence * weight
            confidence_weight += weight

    if current_tokens:
        lines.append(" ".join(current_tokens))

    text = "\n".join(lines).strip()
    confidence = (
        round(weighted_confidence / confidence_weight / 100, 3)
        if confidence_weight
        else None
    )
    return text, confidence


def extract_pages_from_bytes(content: bytes, filename: str) -> list[ExtractedPage]:
    """Постранично извлечь текст, фиксируя способ извлечения.

    Постраничность нужна для прослеживаемости: каждое извлечённое поле
    должно ссылаться на конкретную страницу источника.

    Для PDF сначала читается текстовый слой. OCR включается постранично
    только там, где текста недостаточно; это сохраняет структуру обычных
    PDF и не расходует CPU без необходимости.
    """
    ext = detect_extension_for_pages(filename)

    if ext == ".pdf":
        if not _PDFPLUMBER_AVAILABLE:
            raise RuntimeError("pdfplumber не установлен. Выполните: pip install pdfplumber")
        pages: list[ExtractedPage] = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for index, page in enumerate(pdf.pages, start=1):
                text = page.extract_text() or ""
                if len(text.strip()) < settings.OCR_MIN_TEXT_CHARS:
                    try:
                        image = page.to_image(
                            resolution=settings.OCR_DPI, antialias=True
                        ).original
                        ocr_text, confidence = _ocr_image(image)
                    except NoTextLayerError:
                        # Короткая подпись в цифровом PDF всё равно полезна.
                        # Ошибка OCR критична только для действительно пустой страницы.
                        if text.strip():
                            logger.warning(
                                "OCR fallback failed for PDF page %d; using text layer",
                                index,
                            )
                            ocr_text, confidence = "", None
                        else:
                            raise
                    if ocr_text.strip():
                        pages.append(
                            ExtractedPage(
                                page_number=index,
                                text=ocr_text,
                                method="ocr",
                                ocr_confidence=confidence,
                            )
                        )
                        continue
                pages.append(
                    ExtractedPage(page_number=index, text=text, method="pdf_text_layer")
                )
        if not any(p.text.strip() for p in pages):
            raise NoTextLayerError(
                "Не удалось распознать текст в PDF. Проверьте качество скана."
            )
        return pages

    if ext in (".docx", ".doc"):
        # python-docx не даёт разбиения на страницы: разбивку определяет
        # рендерер Word. Возвращаем один блок и честно помечаем его как
        # страницу 1, а не выдумываем границы страниц.
        return [
            ExtractedPage(page_number=1, text=extract_docx_text(content), method="docx_parser")
        ]

    if ext == ".txt":
        for encoding in ("utf-8", "cp1251"):
            try:
                return [
                    ExtractedPage(page_number=1, text=content.decode(encoding), method="plain_text")
                ]
            except UnicodeDecodeError:
                continue
        raise UnsupportedDocumentType("Не удалось декодировать TXT (пробовали UTF-8 и CP1251)")

    if ext in (".png", ".jpg", ".jpeg"):
        if not _OCR_IMPORTS_AVAILABLE:
            raise NoTextLayerError(
                "OCR недоступен: не установлен Python-пакет pytesseract или Pillow."
            )
        try:
            with Image.open(io.BytesIO(content)) as image:
                text, confidence = _ocr_image(image)
        except NoTextLayerError:
            raise
        except Exception as exc:  # noqa: BLE001
            raise UnsupportedDocumentType(
                f"Не удалось открыть изображение: {exc}"
            ) from exc
        if not text.strip():
            raise NoTextLayerError(
                "Не удалось распознать текст на изображении. Проверьте качество снимка."
            )
        return [
            ExtractedPage(
                page_number=1,
                text=text,
                method="ocr",
                ocr_confidence=confidence,
            )
        ]

    raise UnsupportedDocumentType(f"Неподдерживаемое расширение: {ext}")


class NoTextLayerError(ValueError):
    """В документе нет извлекаемого текста — нужен OCR."""


def detect_extension_for_pages(filename: str) -> str:
    """Как detect_extension, но допускает также изображения."""
    name = (filename or "").lower()
    for ext in (".pdf", ".docx", ".doc", ".txt", ".png", ".jpeg", ".jpg"):
        if name.endswith(ext):
            return ext
    raise UnsupportedDocumentType(f"Неподдерживаемый тип файла: {filename}")


def extract_from_path(path) -> str:
    """Утилита для тестов: читает файл по пути (str или Path)."""
    p = Path(path)
    return extract_text_from_bytes(p.read_bytes(), p.name)
