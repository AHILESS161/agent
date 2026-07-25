"""
DocumentGenerator — builds DOCX documents programmatically for the Trademark
Registration System.

Documents generated:
    1. Application Draft   — Заявление на регистрацию товарного знака
    2. Missing Info Letter — Письмо клиенту о нехватке данных
    3. Legal Memo          — Внутренний меморандум для юриста

All documents are created using python-docx without relying on template files
(MVP approach). The interface is ready for migration to docxtpl Jinja2 templates
in production.

Usage::

    from app.services.document_generator import DocumentGenerator
    gen = DocumentGenerator(output_dir=Path("generated_docs"))

    docx_path = gen.generate_application_draft(application, client)
    letter_path = gen.generate_missing_info_letter(application, ["Изображение знака", "ИНН"])
    memo_path = gen.generate_legal_memo(application, legal_review)
"""

from __future__ import annotations

import io
import logging
from datetime import date, datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)

# Graceful import — python-docx may not be installed in all environments
try:
    from docx import Document as _DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    logger.warning(
        "python-docx not installed. DocumentGenerator will run in stub mode. "
        "Install with: pip install python-docx"
    )


# ---------------------------------------------------------------------------
# Exceptions
# ---------------------------------------------------------------------------

class CompletenessError(ValueError):
    """
    Raised when a required field is missing before document generation.

    Attributes:
        missing_fields: List of field names that are required but absent.
    """

    def __init__(self, missing_fields: list[str]) -> None:
        self.missing_fields = missing_fields
        fields_str = ", ".join(missing_fields)
        super().__init__(
            f"Cannot generate document: required fields are missing: {fields_str}"
        )


# ---------------------------------------------------------------------------
# Formatting helpers
# ---------------------------------------------------------------------------

_MARK_TYPE_LABELS = {
    "word": "Словесное",
    "figurative": "Изобразительное",
    "combined": "Комбинированное",
    "3d": "Объёмное",
    "sound": "Звуковое",
    "color": "Цветовое",
    "other": "Иное",
}

_RISK_LABELS = {
    "low": "Низкий",
    "medium": "Средний",
    "high": "Высокий",
    "critical": "Критический",
}

_CLIENT_TYPE_LABELS = {
    "company": "Юридическое лицо",
    "individual": "Физическое лицо",
    "sole_proprietor": "Индивидуальный предприниматель",
}


def _fmt_date(dt: Optional[datetime | date | str] = None) -> str:
    """Format date as DD.MM.YYYY.

    Accepts ``datetime``, ``date``, ISO-formatted strings (``YYYY-MM-DD``
    or full ISO 8601 with time) and ``None``. Anything unparsable is
    returned as ``""`` rather than raising.
    """
    if dt is None:
        return ""
    if isinstance(dt, str):
        # Try parsing ISO date / datetime string.
        try:
            dt = datetime.fromisoformat(dt.replace("Z", "+00:00"))
        except ValueError:
            try:
                dt = datetime.strptime(dt, "%Y-%m-%d")
            except ValueError:
                return dt
    if isinstance(dt, datetime):
        dt = dt.date()
    if isinstance(dt, date):
        return dt.strftime("%d.%m.%Y")
    return str(dt)


def _today() -> str:
    return _fmt_date()


# ---------------------------------------------------------------------------
# DocumentGenerator
# ---------------------------------------------------------------------------

class DocumentGenerator:
    """
    Generates DOCX documents for trademark registration workflow.

    Args:
        output_dir: Directory where generated files are saved.
                    Created automatically if it doesn't exist.
    """

    def __init__(self, output_dir: Optional[Path] = None) -> None:
        self.output_dir = Path(output_dir or "generated_docs")
        self.output_dir.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    def generate_application_draft(
        self,
        application: object,
        client: object,
        *,
        representative: object | None = None,
        mktu_classes: list[int] | None = None,
        priority_claim_info: dict | None = None,
    ) -> Path:
        """
        Generate the official trademark registration application form (DOCX).

        Implements the structure of the standard Russian Federation application
        form used by Rospatent (Федеральная служба по интеллектуальной
        собственности), covering sections:

            А. Заявитель
            Б. Место нахождения / жительства
            В. Представитель заявителя (патентный поверенный)
            Г. Адрес для переписки
            Д. Заявляемое обозначение
            Е. Товары и услуги / Классы МКТУ
            Ж. Приоритет
            З. Подпись

        Args:
            application: TrademarkApplicationDraft ORM instance (or compatible dict).
            client: Client ORM instance (or compatible dict).
            representative: Optional ClientRepresentative ORM instance — патентный
                поверенный или представитель по доверенности.
            mktu_classes: Список подтверждённых юристом классов МКТУ.
            priority_claim_info: dict с полями:
                - country: str
                - number: str
                - filing_date: str | datetime | date

        Returns:
            Path to the generated .docx file.

        Raises:
            CompletenessError: If required fields are missing.
        """
        app = _as_dict(application)
        cli = _as_dict(client)
        rep = _as_dict(representative) if representative else {}

        # Completeness check — required by Russian form section А/Д/Е
        missing: list[str] = []
        if not cli.get("full_name_or_company_name") and not cli.get("legal_name"):
            missing.append("client.full_name_or_company_name")
        if not cli.get("inn"):
            missing.append("client.inn")
        if not cli.get("address") and not cli.get("legal_address"):
            missing.append("client.address")
        if not app.get("mark_name") and not app.get("mark_text"):
            missing.append("application.mark_name")
        if not app.get("mark_type"):
            missing.append("application.mark_type")
        if not app.get("goods_services_raw") and not app.get("goods_services_description"):
            missing.append("application.goods_services_raw")

        if missing:
            raise CompletenessError(missing)

        if not _DOCX_AVAILABLE:
            return self._stub_file("application_draft", app.get("id", "unknown"))

        doc = _DocxDocument()
        _set_margins(doc)

        # Header
        _add_heading(doc, "ФЕДЕРАЛЬНАЯ СЛУЖБА ПО ИНТЕЛЛЕКТУАЛЬНОЙ СОБСТВЕННОСТИ", 1)
        _add_heading(doc, "(РОСПАТЕНТ)", 2)
        doc.add_paragraph()
        _add_heading(doc, "ЗАЯВКА НА РЕГИСТРАЦИЮ ТОВАРНОГО ЗНАКА", 1)
        doc.add_paragraph()

        # Date and application number
        _add_labeled(doc, "Дата подачи:", _today())
        _add_labeled(doc, "Вх. номер (заполняется Роспатентом):", "________________")

        # ------------------------------------------------------------------
        # А. ЗАЯВИТЕЛЬ
        # ------------------------------------------------------------------
        doc.add_paragraph()
        _add_section_heading(doc, "А. ЗАЯВИТЕЛЬ")
        client_type = cli.get("type") or cli.get("client_type", "company")
        _add_labeled(
            doc,
            "Вид заявителя (юр. лицо / ИП / физ. лицо):",
            _CLIENT_TYPE_LABELS.get(str(client_type), str(client_type)),
        )
        legal_name = (
            cli.get("full_name_or_company_name")
            or cli.get("legal_name")
            or ""
        )
        _add_labeled(doc, "Полное наименование / ФИО:", legal_name)
        if cli.get("short_name"):
            _add_labeled(doc, "Сокращённое наименование:", cli.get("short_name", ""))
        _add_labeled(doc, "ИНН:", cli.get("inn", ""))
        ogrn = cli.get("ogrn_or_ogrnip") or cli.get("ogrn")
        if ogrn:
            _add_labeled(doc, "ОГРН / ОГРНИП:", ogrn)
        if cli.get("country"):
            _add_labeled(doc, "Страна:", cli.get("country", ""))

        # ------------------------------------------------------------------
        # Б. МЕСТО НАХОЖДЕНИЯ / ЖИТЕЛЬСТВА
        # ------------------------------------------------------------------
        doc.add_paragraph()
        _add_section_heading(doc, "Б. МЕСТО НАХОЖДЕНИЯ / ЖИТЕЛЬСТВА")
        _add_labeled(
            doc,
            "Юридический / почтовый адрес (с индексом):",
            cli.get("address") or cli.get("legal_address") or "",
        )

        # ------------------------------------------------------------------
        # В. ПРЕДСТАВИТЕЛЬ ЗАЯВИТЕЛЯ
        # ------------------------------------------------------------------
        doc.add_paragraph()
        _add_section_heading(doc, "В. ПРЕДСТАВИТЕЛЬ ЗАЯВИТЕЛЯ (ПАТЕНТНЫЙ ПОВЕРЕННЫЙ)")
        if rep:
            _add_labeled(doc, "ФИО представителя:", rep.get("full_name", ""))
            if rep.get("role"):
                _add_labeled(doc, "Статус / роль:", rep.get("role", ""))
            if rep.get("poa_reference"):
                _add_labeled(
                    doc,
                    "Доверенность № / дата:",
                    rep.get("poa_reference", ""),
                )
            if rep.get("phone"):
                _add_labeled(doc, "Телефон представителя:", rep.get("phone", ""))
            if rep.get("email"):
                _add_labeled(doc, "Email представителя:", rep.get("email", ""))
            if rep.get("personal_data_consent_reference"):
                _add_labeled(
                    doc,
                    "Согласие на обработку ПДн:",
                    rep.get("personal_data_consent_reference", ""),
                )
        else:
            doc.add_paragraph(
                "(Представитель не указан. Заявитель действует самостоятельно.)"
            )

        # ------------------------------------------------------------------
        # Г. АДРЕС ДЛЯ ПЕРЕПИСКИ
        # ------------------------------------------------------------------
        doc.add_paragraph()
        _add_section_heading(doc, "Г. АДРЕС ДЛЯ ПЕРЕПИСКИ")
        corr_address = (
            rep.get("address")
            if rep and rep.get("address")
            else (cli.get("address") or cli.get("legal_address") or "")
        )
        _add_labeled(doc, "Адрес для переписки:", corr_address)
        _add_labeled(
            doc,
            "Телефон для переписки:",
            (rep.get("phone") if rep and rep.get("phone") else cli.get("phone", "")) or "",
        )
        _add_labeled(
            doc,
            "Электронная почта для переписки:",
            (rep.get("email") if rep and rep.get("email") else cli.get("email", "")) or "",
        )

        # ------------------------------------------------------------------
        # Д. ЗАЯВЛЯЕМОЕ ОБОЗНАЧЕНИЕ
        # ------------------------------------------------------------------
        doc.add_paragraph()
        _add_section_heading(doc, "Д. ЗАЯВЛЯЕМОЕ ОБОЗНАЧЕНИЕ")
        mark_name = app.get("mark_name") or app.get("mark_text") or ""
        _add_labeled(doc, "Обозначение (полное):", mark_name)
        mark_text = app.get("mark_text") or app.get("mark_name") or ""
        _add_labeled(doc, "Словесный элемент:", mark_text)
        mark_type = app.get("mark_type", "word")
        _add_labeled(
            doc,
            "Вид обозначения:",
            _MARK_TYPE_LABELS.get(str(mark_type), str(mark_type)),
        )
        if app.get("colors_claimed"):
            _add_labeled(
                doc,
                "Заявляемые цвета:",
                app.get("colors_claimed", ""),
            )
        if app.get("transliteration"):
            _add_labeled(
                doc,
                "Транслитерация:",
                app.get("transliteration", ""),
            )
        if app.get("translation"):
            _add_labeled(doc, "Перевод:", app.get("translation", ""))
        if app.get("description_of_mark"):
            _add_labeled(
                doc,
                "Описание обозначения:",
                app.get("description_of_mark", ""),
            )
        if app.get("mark_image_file_id"):
            _add_labeled(
                doc,
                "Файл изображения обозначения:",
                app.get("mark_image_file_id", ""),
            )

        # ------------------------------------------------------------------
        # Е. ТОВАРЫ И УСЛУГИ / КЛАССЫ МКТУ
        # ------------------------------------------------------------------
        doc.add_paragraph()
        _add_section_heading(doc, "Е. ТОВАРЫ И УСЛУГИ / КЛАССЫ МКТУ")
        goods_raw = (
            app.get("goods_services_raw")
            or app.get("goods_services_description")
            or ""
        )
        doc.add_paragraph(goods_raw)
        if mktu_classes:
            classes_str = ", ".join(str(c) for c in sorted(set(mktu_classes)))
        else:
            classes_str = "(Классы МКТУ подтверждаются юристом перед подачей)"
        _add_labeled(doc, "Классы МКТУ:", classes_str)

        # ------------------------------------------------------------------
        # Ж. ПРИОРИТЕТ
        # ------------------------------------------------------------------
        if priority_claim_info:
            doc.add_paragraph()
            _add_section_heading(doc, "Ж. ПРИОРИТЕТ")
            if priority_claim_info.get("country"):
                _add_labeled(
                    doc,
                    "Страна подачи первоначальной заявки:",
                    str(priority_claim_info.get("country", "")),
                )
            if priority_claim_info.get("number"):
                _add_labeled(
                    doc,
                    "Номер первоначальной заявки:",
                    str(priority_claim_info.get("number", "")),
                )
            if priority_claim_info.get("filing_date"):
                _add_labeled(
                    doc,
                    "Дата подачи первоначальной заявки:",
                    _fmt_date(priority_claim_info.get("filing_date")),
                )

        # ------------------------------------------------------------------
        # З. ПОДПИСЬ
        # ------------------------------------------------------------------
        doc.add_paragraph()
        _add_section_heading(doc, "З. ПОДПИСЬ")
        doc.add_paragraph()
        signer_label = (
            f"Подпись {rep.get('full_name', '') or legal_name}:"
            if rep and rep.get("full_name")
            else f"Подпись {legal_name}:"
        )
        _add_labeled(doc, signer_label, "____________________")
        _add_labeled(doc, "ФИО подписанта:", legal_name)
        _add_labeled(doc, "Дата:", _today())

        filename = f"application_draft_{app.get('id', 'draft')}_{date.today().isoformat()}.docx"
        out_path = self.output_dir / filename
        doc.save(str(out_path))
        logger.info(f"DocumentGenerator: saved application draft → {out_path}")
        return out_path

    def generate_missing_info_letter(
        self,
        application: object,
        missing_items: list[str],
        deadline_days: int = 14,
    ) -> Path:
        """
        Generate a letter to the client requesting missing information.

        Args:
            application: TrademarkApplicationDraft ORM instance.
            missing_items: List of human-readable missing item descriptions.
            deadline_days: Days from today to include as the response deadline.

        Returns:
            Path to the generated .docx file.

        Raises:
            CompletenessError: If application has no ID or client data.
        """
        app = _as_dict(application)

        if not missing_items:
            raise CompletenessError(["missing_items (список не может быть пустым)"])

        if not _DOCX_AVAILABLE:
            return self._stub_file("missing_info_letter", app.get("id", "unknown"))

        doc = _DocxDocument()
        _set_margins(doc)

        # Sender info
        _add_paragraph_right(
            doc,
            f"Исх. № {app.get('id', 'N/A')}-ЗТЗ\nот {_today()}",
        )
        doc.add_paragraph()

        # Recipient
        client_name = app.get("client_legal_name", "Клиенту")
        _add_paragraph_right(doc, f"Заявителю:\n{client_name}")
        doc.add_paragraph()

        # Subject
        _add_heading(
            doc,
            f"ЗАПРОС ДОПОЛНИТЕЛЬНЫХ СВЕДЕНИЙ\nпо заявке «{app.get('mark_name', '')}»",
            2,
        )
        doc.add_paragraph()

        # Body
        doc.add_paragraph(
            f"Уважаемый(-ая) заявитель,\n\n"
            f"В ходе рассмотрения Вашей заявки на регистрацию товарного знака "
            f"«{app.get('mark_name', '')}» (внутренний номер: {app.get('id', 'N/A')}) "
            f"выявлена необходимость предоставления дополнительных сведений и/или документов."
        )
        doc.add_paragraph()
        doc.add_paragraph(
            "Просим предоставить следующие документы / сведения:"
        )

        # Numbered list of missing items
        for i, item in enumerate(missing_items, start=1):
            p = doc.add_paragraph(f"{i}. {item}", style="List Number")

        from datetime import timedelta
        deadline = date.today() + timedelta(days=deadline_days)
        doc.add_paragraph()
        doc.add_paragraph(
            f"Срок предоставления ответа: {_fmt_date(deadline)} "
            f"({deadline_days} рабочих дней с даты настоящего письма)."
        )
        doc.add_paragraph()
        doc.add_paragraph(
            "При возникновении вопросов, пожалуйста, свяжитесь с Вашим менеджером."
        )
        doc.add_paragraph()

        # Signature block
        doc.add_paragraph("С уважением,")
        doc.add_paragraph()
        _add_labeled(doc, "Менеджер по сопровождению:", "____________________")
        _add_labeled(doc, "ФИО:", "____________________")
        _add_labeled(doc, "Тел.:", "____________________")
        _add_labeled(doc, "Email:", "____________________")

        filename = (
            f"missing_info_letter_{app.get('id', 'draft')}"
            f"_{date.today().isoformat()}.docx"
        )
        out_path = self.output_dir / filename
        doc.save(str(out_path))
        logger.info(f"DocumentGenerator: saved missing info letter → {out_path}")
        return out_path

    def generate_legal_memo(
        self,
        application: object,
        legal_review: object,
    ) -> Path:
        """
        Generate an internal legal memo for the lawyer.

        Args:
            application: TrademarkApplicationDraft ORM instance.
            legal_review: LegalReview ORM instance.

        Returns:
            Path to the generated .docx file.

        Raises:
            CompletenessError: If required review data is missing.
        """
        app = _as_dict(application)
        review = _as_dict(legal_review)

        missing: list[str] = []
        if not review.get("summary"):
            missing.append("legal_review.summary")
        if not review.get("risk_level"):
            missing.append("legal_review.risk_level")
        if missing:
            raise CompletenessError(missing)

        if not _DOCX_AVAILABLE:
            return self._stub_file("legal_memo", app.get("id", "unknown"))

        doc = _DocxDocument()
        _set_margins(doc)

        # Header
        _add_heading(doc, "ВНУТРЕННИЙ МЕМОРАНДУМ", 1)
        _add_heading(doc, "Правовая экспертиза товарного знака", 2)
        doc.add_paragraph()

        # Meta block
        table = doc.add_table(rows=6, cols=2)
        table.style = "Table Grid"
        meta_rows = [
            ("Документ:", "Внутренний меморандум"),
            ("Дата:", _today()),
            ("Заявка №:", str(app.get("id", "N/A"))),
            ("Товарный знак:", app.get("mark_name", "")),
            (
                "Уровень риска:",
                _RISK_LABELS.get(str(review.get("risk_level", "")), ""),
            ),
            (
                "Решение:",
                _decision_label(str(review.get("decision", ""))),
            ),
        ]
        for i, (label, value) in enumerate(meta_rows):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            _bold_cell(row.cells[0])

        doc.add_paragraph()
        _add_section_heading(doc, "1. СВЕДЕНИЯ О ЗАЯВКЕ")
        _add_labeled(doc, "Обозначение:", app.get("mark_name", ""))
        mark_type = app.get("mark_type", "word")
        _add_labeled(
            doc,
            "Вид обозначения:",
            _MARK_TYPE_LABELS.get(str(mark_type), str(mark_type)),
        )
        if app.get("goods_services_description"):
            _add_labeled(
                doc,
                "Товары/услуги:",
                app.get("goods_services_description", ""),
            )

        doc.add_paragraph()
        _add_section_heading(doc, "2. РЕЗУЛЬТАТЫ ПРАВОВОЙ ЭКСПЕРТИЗЫ")
        doc.add_paragraph(review.get("summary", ""))

        # Blocking issues
        blocking = review.get("blocking_issues_json") or []
        if blocking:
            doc.add_paragraph()
            _add_section_heading(doc, "3. БЛОКИРУЮЩИЕ НАРУШЕНИЯ")
            for issue in blocking:
                if isinstance(issue, dict):
                    p = doc.add_paragraph(
                        f"• [{issue.get('severity', '').upper()}] "
                        f"{issue.get('message', '')}"
                    )
                    run = p.runs[0]
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
                else:
                    doc.add_paragraph(f"• {issue}")

        # Non-blocking issues
        non_blocking = review.get("non_blocking_issues_json") or []
        if non_blocking:
            doc.add_paragraph()
            heading_num = 4 if blocking else 3
            _add_section_heading(doc, f"{heading_num}. ЗАМЕЧАНИЯ (НЕ БЛОКИРУЮЩИЕ)")
            for issue in non_blocking:
                if isinstance(issue, dict):
                    doc.add_paragraph(
                        f"• {issue.get('message', '')} — "
                        f"{issue.get('recommended_action', '')}"
                    )
                else:
                    doc.add_paragraph(f"• {issue}")

        doc.add_paragraph()
        _add_section_heading(doc, "РЕКОМЕНДАЦИЯ")
        decision = str(review.get("decision", ""))
        risk = str(review.get("risk_level", ""))
        doc.add_paragraph(
            f"На основании проведённого анализа рекомендуется: "
            f"{_decision_label(decision).upper()}. "
            f"Уровень риска: {_RISK_LABELS.get(risk, risk).upper()}."
        )
        doc.add_paragraph()

        # Signature block
        doc.add_paragraph()
        _add_labeled(doc, "Юрист:", "____________________")
        _add_labeled(doc, "ФИО:", "____________________")
        _add_labeled(doc, "Дата:", _today())

        filename = (
            f"legal_memo_{app.get('id', 'draft')}"
            f"_{date.today().isoformat()}.docx"
        )
        out_path = self.output_dir / filename
        doc.save(str(out_path))
        logger.info(f"DocumentGenerator: saved legal memo → {out_path}")
        return out_path

    # ------------------------------------------------------------------
    # Stub fallback (when python-docx not available)
    # ------------------------------------------------------------------

    def _stub_file(self, doc_type: str, app_id: object) -> Path:
        """Return a path to a stub text file when python-docx is unavailable."""
        filename = f"{doc_type}_{app_id}_{date.today().isoformat()}.txt"
        out_path = self.output_dir / filename
        out_path.write_text(
            f"[STUB] {doc_type} for application {app_id}. "
            "Install python-docx to generate real DOCX documents."
        )
        logger.warning(f"DocumentGenerator: stub file created → {out_path}")
        return out_path


# ---------------------------------------------------------------------------
# python-docx helpers
# ---------------------------------------------------------------------------

def _as_dict(obj: object) -> dict:
    """Convert an ORM instance or dict to a plain dict."""
    if isinstance(obj, dict):
        return obj
    # SQLAlchemy ORM objects
    if hasattr(obj, "__dict__"):
        return {k: v for k, v in obj.__dict__.items() if not k.startswith("_")}
    return {}


def _set_margins(doc: "_DocxDocument") -> None:
    """Set standard Russian document margins (30mm left, 20mm others)."""
    if not _DOCX_AVAILABLE:
        return
    from docx.shared import Cm
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)


def _add_heading(doc: "_DocxDocument", text: str, level: int = 1) -> None:
    if not _DOCX_AVAILABLE:
        return
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_section_heading(doc: "_DocxDocument", text: str) -> None:
    if not _DOCX_AVAILABLE:
        return
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)


def _add_labeled(doc: "_DocxDocument", label: str, value: str) -> None:
    if not _DOCX_AVAILABLE:
        return
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label} ")
    run_label.bold = True
    p.add_run(value)


def _add_paragraph_right(doc: "_DocxDocument", text: str) -> None:
    if not _DOCX_AVAILABLE:
        return
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _bold_cell(cell: object) -> None:
    if not _DOCX_AVAILABLE:
        return
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True


def _decision_label(decision: str) -> str:
    labels = {
        "approve": "Подать заявку",
        "reject": "Отказать в регистрации",
        "modify": "Доработать обозначение",
        "further_review": "Дополнительная экспертиза",
    }
    return labels.get(decision, decision)
