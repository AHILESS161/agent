"""Тесты чернового заявления.

Главное правило: в документ попадают только подтверждённые значения,
а экспорт возможен лишь после утверждения специалистом. Черновик
юридически значимого документа не должен содержать непроверенных
данных и не должен уходить наружу как готовый.
"""

from __future__ import annotations

import pytest

from app.infrastructure.database.models import MarkType, UserRole
from app.services import file_storage
from tests.conftest import login_headers

EGRUL_TEXT = (
    "ВЫПИСКА\n"
    "из Единого государственного реестра юридических лиц\n"
    "1 Полное наименование ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ ПРИМЕР\n"
    "13 ОГРН 1027700132195\n"
    "19 ИНН 7707083893\n"
    "Сведения о регистрирующем органе\n"
).encode("utf-8")


@pytest.fixture(autouse=True)
def isolated_storage(tmp_path, monkeypatch):
    monkeypatch.setattr(
        file_storage.settings, "FILE_STORAGE_PATH", str(tmp_path / "docs")
    )
    return tmp_path


@pytest.fixture
async def lawyer_auth(client, api_user_factory) -> dict[str, str]:
    await api_user_factory("lawyer-draft@test.ru", UserRole.lawyer)
    return login_headers(client, "lawyer-draft@test.ru")


@pytest.fixture
def case_with_fields(client, lawyer_auth) -> int:
    """Дело с загруженной выпиской и извлечёнными полями."""
    client_id = client.post(
        "/api/v1/clients",
        json={"type": "company", "full_name_or_company_name": 'ООО "Тест"'},
        headers=lawyer_auth,
    ).json()["id"]
    app_id = client.post(
        "/api/v1/applications",
        json={"client_id": client_id, "mark_name": "ТЕСТЗНАК"},
        headers=lawyer_auth,
    ).json()["id"]

    document = client.post(
        f"/api/v1/applications/{app_id}/source-documents",
        files={"file": ("выписка.txt", EGRUL_TEXT, "text/plain")},
        headers=lawyer_auth,
    ).json()
    client.post(
        f"/api/v1/source-documents/{document['id']}/extract", headers=lawyer_auth
    )
    return app_id


@pytest.mark.api
class TestDraftOnlyUsesConfirmedData:
    """Ключевое требование к черновику."""

    def test_unconfirmed_fields_are_not_filled(
        self, client, lawyer_auth, case_with_fields
    ):
        draft = client.post(
            f"/api/v1/applications/{case_with_fields}/draft", headers=lawyer_auth
        ).json()

        filled_labels = {item["label"] for item in draft["filled_fields"]}
        assert "ОГРН" not in filled_labels
        assert "ИНН" not in filled_labels

    def test_skipped_fields_explain_why(self, client, lawyer_auth, case_with_fields):
        draft = client.post(
            f"/api/v1/applications/{case_with_fields}/draft", headers=lawyer_auth
        ).json()

        assert draft["skipped_fields"]
        for item in draft["skipped_fields"]:
            assert item["reason"]

    def test_confirmed_field_appears_in_draft(
        self, client, lawyer_auth, case_with_fields
    ):
        fields = client.get(
            f"/api/v1/applications/{case_with_fields}/field-reconciliation",
            headers=lawyer_auth,
        ).json()
        ogrn = next(
            item for item in fields["items"] if item["label"] == "ОГРН"
        )
        client.post(
            f"/api/v1/extracted-fields/{ogrn['extracted_field_id']}/confirm",
            json={"action": "accept"},
            headers=lawyer_auth,
        )

        draft = client.post(
            f"/api/v1/applications/{case_with_fields}/draft", headers=lawyer_auth
        ).json()
        filled = {item["label"]: item["value"] for item in draft["filled_fields"]}
        assert filled.get("ОГРН") == "1027700132195"

    def test_filled_field_keeps_its_source(
        self, client, lawyer_auth, case_with_fields
    ):
        """Специалист должен видеть, откуда значение попало в документ."""
        fields = client.get(
            f"/api/v1/applications/{case_with_fields}/field-reconciliation",
            headers=lawyer_auth,
        ).json()
        inn = next(item for item in fields["items"] if item["label"] == "ИНН")
        client.post(
            f"/api/v1/extracted-fields/{inn['extracted_field_id']}/confirm",
            json={"action": "accept"},
            headers=lawyer_auth,
        )

        draft = client.post(
            f"/api/v1/applications/{case_with_fields}/draft", headers=lawyer_auth
        ).json()
        entry = next(i for i in draft["filled_fields"] if i["label"] == "ИНН")
        assert "regex" in entry["source"]


@pytest.mark.api
class TestVersioningAndProvenance:
    def test_each_generation_creates_new_version(
        self, client, lawyer_auth, case_with_fields
    ):
        first = client.post(
            f"/api/v1/applications/{case_with_fields}/draft", headers=lawyer_auth
        ).json()
        second = client.post(
            f"/api/v1/applications/{case_with_fields}/draft", headers=lawyer_auth
        ).json()

        assert second["version"] == first["version"] + 1

    def test_versions_are_listed(self, client, lawyer_auth, case_with_fields):
        client.post(
            f"/api/v1/applications/{case_with_fields}/draft", headers=lawyer_auth
        )
        client.post(
            f"/api/v1/applications/{case_with_fields}/draft", headers=lawyer_auth
        )

        listing = client.get(
            f"/api/v1/applications/{case_with_fields}/drafts", headers=lawyer_auth
        ).json()
        assert listing["total"] == 2

    def test_draft_records_template_and_mapping_versions(
        self, client, lawyer_auth, case_with_fields
    ):
        """Нужно знать, по какому бланку и маппингу собран документ."""
        draft = client.post(
            f"/api/v1/applications/{case_with_fields}/draft", headers=lawyer_auth
        ).json()

        provenance = draft["provenance"]
        assert provenance["template_name"]
        assert provenance["schema_version"]
        assert provenance["mapping_version"] >= 1


@pytest.mark.api
class TestExportRequiresApproval:
    def test_export_blocked_before_approval(
        self, client, lawyer_auth, case_with_fields
    ):
        draft = client.post(
            f"/api/v1/applications/{case_with_fields}/draft", headers=lawyer_auth
        ).json()
        assert draft["can_export"] is False

        response = client.get(f"/api/v1/drafts/{draft['id']}/download", headers=lawyer_auth)
        assert response.status_code == 409
        assert "не утверждён" in response.json()["detail"]

    def test_export_allowed_after_approval(
        self, client, lawyer_auth, case_with_fields
    ):
        draft = client.post(
            f"/api/v1/applications/{case_with_fields}/draft", headers=lawyer_auth
        ).json()
        approved = client.post(
            f"/api/v1/drafts/{draft['id']}/approve", headers=lawyer_auth
        ).json()
        assert approved["status"] == "approved_by_specialist"
        assert approved["can_export"] is True

        response = client.get(f"/api/v1/drafts/{draft['id']}/download", headers=lawyer_auth)
        assert response.status_code == 200
        assert len(response.content) > 1000

    def test_download_marks_draft_exported(
        self, client, lawyer_auth, case_with_fields
    ):
        draft = client.post(
            f"/api/v1/applications/{case_with_fields}/draft", headers=lawyer_auth
        ).json()
        client.post(f"/api/v1/drafts/{draft['id']}/approve", headers=lawyer_auth)
        client.get(f"/api/v1/drafts/{draft['id']}/download", headers=lawyer_auth)

        listing = client.get(
            f"/api/v1/applications/{case_with_fields}/drafts", headers=lawyer_auth
        ).json()
        assert listing["items"][0]["status"] == "exported"

    async def test_manager_cannot_approve(
        self, client, api_user_factory, lawyer_auth, case_with_fields
    ):
        """Утверждение содержания документа — решение специалиста."""
        draft = client.post(
            f"/api/v1/applications/{case_with_fields}/draft", headers=lawyer_auth
        ).json()

        await api_user_factory("manager-draft@test.ru", UserRole.manager)
        manager = login_headers(client, "manager-draft@test.ru")

        response = client.post(f"/api/v1/drafts/{draft['id']}/approve", headers=manager)
        assert response.status_code == 403


@pytest.mark.api
class TestOfficialBlank:
    """Черновик — это заполненный бланк Роспатента, а не своя форма."""

    def _render(self, filled, classes=None, mark_image=None, mark_type="word", include_goods_attachment=False, paper_certificate=False, applicant_type=None):
        from app.services.application_draft import (
            DraftContent,
            FilledField,
            render_docx,
        )

        content = DraftContent(applicant_type=applicant_type)
        content.filled = [
            FilledField(field_id=fid, label=fid, value=value, source="regex")
            for fid, value in filled.items()
        ]
        content.classes = classes or []

        class _App:
            id = 1
            mark_text = "ЗВЁЗДОЧКА"
            mark_name = "ЗВЁЗДОЧКА"
            mark_type = None
            filing_method = "electronic"
            signatory_name = None
            signatory_position = None
            signature_date = None
            request_paper_certificate = False

        _App.mark_type = MarkType(mark_type)
        _App.request_paper_certificate = paper_certificate
        return render_docx(
            content,
            _App(),
            mark_image=mark_image,
            include_goods_attachment=include_goods_attachment,
        )

    def _document(self, payload):
        import io

        import docx

        return docx.Document(io.BytesIO(payload))

    def test_structure_of_the_blank_is_preserved(self):
        """Бланк остаётся бланком: та же таблица, те же строки."""
        document = self._document(self._render({}))
        assert len(document.tables) == 1
        assert len(document.tables[0].rows) == 118

    def test_applicant_lands_in_field_731(self):
        payload = self._render(
            {
                "application.applicant.name": "ООО «ПРИМЕР»",
                "application.applicant.address": "101000, Москва",
            }
        )
        text = self._document(payload).tables[0].rows[5].cells[0].text
        assert "ООО «ПРИМЕР»" in text
        assert "101000, Москва" in text

    def test_identifiers_land_in_their_block(self):
        payload = self._render(
            {
                "application.applicant.ogrn": "1027700132195",
                "application.applicant.inn": "7707083893",
            }
        )
        text = self._document(payload).tables[0].rows[5].cells[31].text
        assert "ОГРН: 1027700132195" in text
        assert "ИНН: 7707083893" in text

    @pytest.mark.parametrize(
        ("applicant_type", "filled", "expected", "forbidden"),
        [
            (
                "company",
                {
                    "application.applicant.ogrn": "1027700132195",
                    "application.applicant.inn": "7707083893",
                    "application.applicant.kpp": "770701001",
                },
                ("ОГРН: 1027700132195", "ИНН: 7707083893", "КПП: 770701001"),
                ("ОГРНИП: 1027700132195",),
            ),
            (
                "sole_proprietor",
                {
                    "application.applicant.ogrn": "315774600312340",
                    "application.applicant.inn": "771234567859",
                    "application.applicant.kpp": "770701001",
                },
                ("ОГРНИП: 315774600312340", "ИНН: 771234567859"),
                ("ОГРН: 315774600312340", "КПП: 770701001"),
            ),
            (
                "individual",
                {
                    "application.applicant.ogrn": "1027700132195",
                    "application.applicant.inn": "771234567859",
                    "application.applicant.kpp": "770701001",
                },
                ("ИНН: 771234567859",),
                (
                    "ОГРН: 1027700132195",
                    "ОГРНИП: 1027700132195",
                    "КПП: 770701001",
                ),
            ),
        ],
    )
    def test_identifier_block_depends_on_applicant_type(
        self, applicant_type, filled, expected, forbidden
    ):
        document = self._document(
            self._render(filled, applicant_type=applicant_type)
        )
        text = document.tables[0].rows[5].cells[31].text
        for value in expected:
            assert value in text
        for value in forbidden:
            assert value not in text

    @pytest.mark.parametrize(
        ("mark_type", "row", "cell"),
        [
            ("word", 26, 15),
            ("figurative", 26, 26),
            ("combined", 34, 1),
            ("sound", 30, 23),
        ],
    )
    def test_mark_type_checkbox_matrix(self, mark_type, row, cell):
        document = self._document(self._render({}, mark_type=mark_type))
        assert "X" in document.tables[0].rows[row].cells[cell].text

    def test_mark_lands_in_field_540(self):
        payload = self._render({"application.mark.text": "ЗВЁЗДОЧКА"})
        document = self._document(payload)
        cell = document.tables[0].rows[17].cells[2]
        assert len(cell.tables) == 1
        mark_cell = cell.tables[0].rows[0].cells[0]
        text = mark_cell.text
        assert "ЗВЁЗДОЧКА" in text
        mark_runs = [
            run
            for paragraph in mark_cell.paragraphs
            for run in paragraph.runs
            if "ЗВЁЗДОЧКА" in run.text
        ]
        assert mark_runs
        assert any(run.bold and run.font.size and run.font.size.pt >= 30 for run in mark_runs)

    def test_image_lands_in_dedicated_box_below_540_and_description_in_571(self):
        import io

        from PIL import Image

        image = io.BytesIO()
        Image.new("RGB", (800, 400), "navy").save(image, format="JPEG")
        payload = self._render(
            {
                "application.mark.text": "ЗВЁЗДОЧКА",
                "application.mark.description": "Комбинированное обозначение",
                "application.mark.transliteration": "ZVEZDOCHKA",
                "application.mark.translation": "STAR",
            },
            mark_image=image.getvalue(),
            mark_type="combined",
        )
        document = self._document(payload)
        outer = document.tables[0].rows[17].cells[2]
        left, right = outer.tables[0].rows[0].cells
        image_box = document.tables[0].rows[19].cells[1]
        assert not left._tc.xpath(".//w:drawing")
        assert "ЗВЁЗДОЧКА" not in left.text
        assert image_box._tc.xpath(".//w:drawing")
        assert "Комбинированное обозначение" in right.text
        assert "Транслитерация" not in right.text
        assert "Перевод" not in right.text
        assert not right._tc.xpath(".//w:drawing")

    def test_paper_certificate_checkbox_follows_user_choice(self):
        unchecked = self._document(self._render({})).tables[0]
        checked = self._document(
            self._render({}, paper_certificate=True)
        ).tables[0]
        assert "X" not in unchecked.rows[83].cells[1].text
        assert "X" in checked.rows[83].cells[1].text

    def test_known_checkboxes_are_filled_automatically(self):
        import io

        from PIL import Image

        image = io.BytesIO()
        Image.new("RGB", (800, 400), "navy").save(image, format="JPEG")
        document = self._document(self._render(
            {"application.mark.colors": "тёмно-синий"},
            classes=[("25", "Одежда")],
            mark_image=image.getvalue(),
            mark_type="combined",
            include_goods_attachment=True,
        ))
        table = document.tables[0]
        assert "X" in table.rows[34].cells[1].text  # комбинированный знак
        assert "X" in table.rows[23].cells[1].text  # цветное исполнение
        assert "X" in table.rows[90].cells[1].text  # изображение приложено
        assert "X" in table.rows[93].cells[1].text  # отдельный перечень товаров
        assert "1" in table.rows[90].cells[37].text
        assert "1" in table.rows[90].cells[43].text
        assert "1" in table.rows[93].cells[37].text
        assert "1" in table.rows[93].cells[43].text

    def test_signatory_details_are_filled_without_fake_drawn_signature(self):
        from datetime import date

        payload = self._render({})
        application = self._document(payload)
        assert not application.tables[0].rows[116].cells[0].text.strip()

        from app.services.application_draft import DraftContent, render_docx

        class _SignedApp:
            id = 2
            mark_text = "ЗВЁЗДОЧКА"
            mark_name = "ЗВЁЗДОЧКА"
            mark_type = MarkType.word
            filing_method = "paper"
            signatory_name = "Иванов Иван Иванович"
            signatory_position = "генеральный директор"
            signature_date = date(2026, 8, 26)

        signed = self._document(render_docx(DraftContent(), _SignedApp()))
        signature_text = signed.tables[0].rows[116].cells[0].text
        date_text = signed.tables[0].rows[117].cells[0].text
        assert "________________" in signature_text
        assert "Иванов Иван Иванович" in signature_text
        assert "генеральный директор" in signature_text
        assert "26.08.2026" in date_text

    def test_color_claim_conflicting_with_monochrome_image_is_not_filed(self):
        import io

        from PIL import Image

        image = io.BytesIO()
        Image.new("RGB", (800, 400), "black").save(image, format="JPEG")
        document = self._document(self._render(
            {"application.mark.colors": "КРАСНЫЙ"},
            mark_image=image.getvalue(),
            mark_type="combined",
        ))
        table = document.tables[0]
        assert "X" not in table.rows[23].cells[1].text
        assert "КРАСНЫЙ" not in table.rows[23].cells[2].text

    def test_publication_caption_and_headers_are_removed(self):
        document = self._document(self._render({"application.mark.text": "ЗВЁЗДОЧКА"}))
        body_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
        assert "Приложение № 1" not in body_text
        for section in document.sections:
            assert not "".join(p.text for p in section.header.paragraphs).strip()
            assert not "".join(p.text for p in section.footer.paragraphs).strip()

    def test_classes_fill_the_goods_table(self):
        payload = self._render({}, classes=[("25", "Одежда; обувь")])
        rows = self._document(payload).tables[0].rows
        assert "25" in rows[44].cells[0].text
        assert "Одежда" in rows[44].cells[3].text

    def test_no_service_notes_inside_the_document(self):
        """В бланке заявления не должно быть пометок про черновик и AI."""
        payload = self._render({"application.mark.text": "ЗВЁЗДОЧКА"})
        document = self._document(payload)
        full = "\n".join(p.text for p in document.paragraphs)
        full += "\n".join(
            cell.text for row in document.tables[0].rows for cell in row.cells
        )

        for forbidden in ("ЧЕРНОВИК", " AI ", "искусственн", "автоматически"):
            assert forbidden not in full, forbidden

    def test_unconfirmed_fields_stay_empty(self):
        """Незаполненное поле остаётся пустым, как в бумажной форме."""
        document = self._document(self._render({}))
        identifiers = document.tables[0].rows[5].cells[31].text
        assert "ОГРН:" in identifiers
        assert "1027700132195" not in identifiers


@pytest.mark.api
class TestChecklist:
    def test_checklist_lists_manual_steps(self, client, lawyer_auth, case_with_fields):
        """Чек-лист объясняет реальный способ подписания, а не старые ручные поля."""
        draft = client.post(
            f"/api/v1/applications/{case_with_fields}/draft", headers=lawyer_auth
        ).json()

        text = " ".join(draft["checklist"])
        assert "ФИО человека, который подпишет" in text
        assert "собственноручную подпись" in text
        assert "электронной подписью" in text

    def test_checklist_mentions_missing_image(
        self, client, lawyer_auth, case_with_fields
    ):
        draft = client.post(
            f"/api/v1/applications/{case_with_fields}/draft", headers=lawyer_auth
        ).json()
        assert any("540" in item for item in draft["checklist"])


@pytest.mark.api
class TestAuth:
    def test_generation_requires_auth(self, client, case_with_fields):
        assert (
            client.post(f"/api/v1/applications/{case_with_fields}/draft").status_code
            == 401
        )

    def test_download_requires_auth(self, client):
        assert client.get("/api/v1/drafts/1/download").status_code == 401
