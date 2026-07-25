"""Unit tests for DocumentGenerator — application_draft section coverage."""

from __future__ import annotations

import tempfile
from pathlib import Path
from typing import Any

import pytest

pytest.importorskip("docx", reason="python-docx is required for these tests")

from docx import Document  # noqa: E402

from app.services.document_generator import (  # noqa: E402
    CompletenessError,
    DocumentGenerator,
)


# ---------------------------------------------------------------------------
# Fixtures / helpers
# ---------------------------------------------------------------------------


def _make_client(**overrides: Any) -> dict[str, Any]:
    base = {
        "id": 1,
        "type": "company",
        "full_name_or_company_name": 'ООО "Ромашка"',
        "short_name": 'ООО "Ромашка"',
        "inn": "7701234567",
        "ogrn_or_ogrnip": "1027701234567",
        "address": "125009, г. Москва, ул. Тверская, д. 1, оф. 100",
        "country": "Российская Федерация",
        "email": "info@romashka.ru",
        "phone": "+7 (495) 123-45-67",
    }
    base.update(overrides)
    return base


def _make_application(**overrides: Any) -> dict[str, Any]:
    base = {
        "id": 42,
        "mark_type": "combined",
        "mark_name": "Ромашка Плюс",
        "mark_text": "Ромашка Плюс",
        "mark_image_file_id": "uploads/marks/42.png",
        "colors_claimed": "зелёный, белый",
        "transliteration": "Romashka Plus",
        "translation": "Daisy Plus",
        "description_of_mark": "Стилизованное изображение цветка ромашки с надписью",
        "goods_services_raw": "программное обеспечение; услуги по обучению",
        "priority_claim": "DE, № 30 2024 000 001 от 12.01.2024",
        "notes": None,
    }
    base.update(overrides)
    return base


def _make_representative(**overrides: Any) -> dict[str, Any]:
    base = {
        "id": 5,
        "client_id": 1,
        "full_name": "Иванов Иван Иванович",
        "email": "i.ivanov@romashka.ru",
        "phone": "+7 (495) 765-43-21",
        "role": "Патентный поверенный",
        "poa_reference": "Доверенность № 17 от 01.03.2024",
        "personal_data_consent_reference": "Согласие № 42 от 01.03.2024",
    }
    base.update(overrides)
    return base


@pytest.fixture
def generator(tmp_path: Path) -> DocumentGenerator:
    return DocumentGenerator(output_dir=tmp_path)


def _all_text(doc: Document) -> str:
    return "\n".join(p.text for p in doc.paragraphs)


# ---------------------------------------------------------------------------
# Completeness
# ---------------------------------------------------------------------------


def test_application_draft_blocks_when_required_fields_missing(
    generator: DocumentGenerator, tmp_path: Path
) -> None:
    app = _make_application()
    bad_client = _make_client(inn=None)
    with pytest.raises(CompletenessError) as exc:
        generator.generate_application_draft(app, bad_client)
    assert "client.inn" in exc.value.missing_fields


def test_application_draft_blocks_when_mark_name_missing(
    generator: DocumentGenerator,
) -> None:
    app = _make_application(mark_name=None, mark_text=None)
    cli = _make_client()
    with pytest.raises(CompletenessError):
        generator.generate_application_draft(app, cli)


# ---------------------------------------------------------------------------
# Section coverage
# ---------------------------------------------------------------------------


def test_application_draft_contains_all_required_sections(
    generator: DocumentGenerator,
) -> None:
    out = generator.generate_application_draft(
        _make_application(),
        _make_client(),
        representative=_make_representative(),
        mktu_classes=[9, 42],
        priority_claim_info={
            "country": "DE",
            "number": "30 2024 000 001",
            "filing_date": "2024-01-12",
        },
    )
    assert out.exists()
    text = _all_text(Document(str(out)))

    # Header
    assert "ФЕДЕРАЛЬНАЯ СЛУЖБА ПО ИНТЕЛЛЕКТУАЛЬНОЙ СОБСТВЕННОСТИ" in text
    assert "ЗАЯВКА НА РЕГИСТРАЦИЮ ТОВАРНОГО ЗНАКА" in text

    # Section А
    assert 'ООО "Ромашка"' in text
    assert "7701234567" in text
    assert "1027701234567" in text

    # Section Б
    assert "г. Москва" in text

    # Section В
    assert "Иванов Иван Иванович" in text
    assert "Патентный поверенный" in text
    assert "Доверенность № 17" in text

    # Section Г
    assert "Адрес для переписки" in text
    assert "i.ivanov@romashka.ru" in text

    # Section Д
    assert "Комбинированное" in text
    assert "Транслитерация" in text
    assert "Romashka Plus" in text
    assert "зелёный, белый" in text
    assert "программное обеспечение" in text

    # Section Е
    assert "9, 42" in text

    # Section Ж
    assert "DE" in text
    assert "30 2024 000 001" in text
    assert "12.01.2024" in text

    # Section З
    assert "Подпись" in text


def test_application_draft_works_without_representative(
    generator: DocumentGenerator,
) -> None:
    out = generator.generate_application_draft(
        _make_application(),
        _make_client(),
    )
    text = _all_text(Document(str(out)))
    assert "Представитель не указан" in text
    assert "info@romashka.ru" in text  # client email used as correspondence email


def test_application_draft_omits_priority_section_when_not_claimed(
    generator: DocumentGenerator,
) -> None:
    out = generator.generate_application_draft(
        _make_application(),
        _make_client(),
    )
    text = _all_text(Document(str(out)))
    assert "Ж. ПРИОРИТЕТ" not in text