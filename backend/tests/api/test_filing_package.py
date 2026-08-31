"""Полный ZIP-пакет для самостоятельной подачи."""

from __future__ import annotations

import io
import zipfile
from datetime import date

import docx
import pytest
from sqlalchemy.ext.asyncio import AsyncSession, async_sessionmaker

from app.infrastructure.database.models import (
    AnalysisKind,
    ApplicationStatus,
    AuditLog,
    Client,
    ClientType,
    ExtractedField,
    ExtractionMethod,
    FieldStatus,
    MarkType,
    NiceCategory,
    NiceClassSuggestion,
    RiskAssessment,
    RiskLevel,
    SearchMode,
    TrademarkApplicationDraft,
    UserRole,
)
from tests.conftest import login_headers


@pytest.mark.api
class TestFilingPackage:
    async def test_application_uses_only_selected_representative_and_fills_draft(
        self, client, api_user_factory
    ):
        await api_user_factory("representative@test.ru", UserRole.client)
        headers = login_headers(client, "representative@test.ru")
        applicant = client.post(
            "/api/v1/clients",
            json={
                "type": "individual",
                "full_name_or_company_name": "Иван Тестов",
                "address": "123456, Москва, ул. Тестовая, д. 1",
            },
            headers=headers,
        ).json()
        other_applicant = client.post(
            "/api/v1/clients",
            json={"type": "individual", "full_name_or_company_name": "Другой заявитель"},
            headers=headers,
        ).json()
        application = client.post(
            "/api/v1/applications",
            json={"client_id": applicant["id"], "mark_name": "ТЕСТ", "mark_type": "word"},
            headers=headers,
        ).json()
        representative = client.post(
            f"/api/v1/clients/{applicant['id']}/representatives",
            json={
                "full_name": "Петров Пётр Петрович",
                "address": "101000, Москва, ул. Представителя, д. 2",
                "email": "patent@example.com",
                "is_patent_attorney": True,
                "patent_attorney_registration_number": "1234",
                "authority_type": "power_of_attorney",
                "poa_reference": "№ 12 от 28.08.2026",
            },
            headers=headers,
        ).json()
        foreign_representative = client.post(
            f"/api/v1/clients/{other_applicant['id']}/representatives",
            json={"full_name": "Чужой представитель"},
            headers=headers,
        ).json()

        rejected = client.put(
            f"/api/v1/applications/{application['id']}",
            json={"representative_id": foreign_representative["id"]},
            headers=headers,
        )
        assert rejected.status_code == 422

        selected = client.put(
            f"/api/v1/applications/{application['id']}",
            json={"representative_id": representative["id"]},
            headers=headers,
        )
        assert selected.status_code == 200, selected.text
        assert selected.json()["representative_id"] == representative["id"]

        package = client.get(
            f"/api/v1/applications/{application['id']}/filing-package",
            headers=headers,
        ).json()
        rules = {item["code"]: item for item in package["requirements"]["requirements"]}
        assert rules["representative_name"]["satisfied"] is True
        assert rules["patent_attorney_number"]["satisfied"] is True
        assert rules["power_of_attorney"]["required"] is True
        assert rules["power_of_attorney"]["satisfied"] is False

        upload = client.post(
            f"/api/v1/applications/{application['id']}/source-documents",
            files={"file": ("doverennost.txt", "Доверенность № 12", "text/plain")},
            headers=headers,
        )
        assert upload.status_code == 201, upload.text
        confirmed = client.put(
            f"/api/v1/source-documents/{upload.json()['id']}/kind",
            json={"document_kind": "power_of_attorney"},
            headers=headers,
        )
        assert confirmed.status_code == 200, confirmed.text

        refreshed = client.get(
            f"/api/v1/applications/{application['id']}/filing-package",
            headers=headers,
        ).json()
        rules = {item["code"]: item for item in refreshed["requirements"]["requirements"]}
        assert rules["power_of_attorney"]["satisfied"] is True

        draft = client.get(
            f"/api/v1/applications/{application['id']}/draft-preview/download",
            headers=headers,
        )
        assert draft.status_code == 200, draft.text
        document = docx.Document(io.BytesIO(draft.content))
        text = "\n".join(
            cell.text for table in document.tables for row in table.rows for cell in row.cells
        )
        assert "Петров Пётр Петрович" in text
        assert "101000, Москва, ул. Представителя, д. 2" in text
        assert "№ 12 от 28.08.2026" in text
        form = document.tables[0]
        assert "X" in form.rows[9].cells[34].text
        assert "X" not in form.rows[11].cells[34].text

        cleared = client.put(
            f"/api/v1/applications/{application['id']}",
            json={"representative_id": None},
            headers=headers,
        )
        assert cleared.status_code == 200, cleared.text
        assert cleared.json()["representative_id"] is None

    async def test_field_sources_are_verified_instead_of_guessed(
        self, client, api_user_factory, async_engine
    ):
        await api_user_factory("filing-sources@test.ru", UserRole.client)
        headers = login_headers(client, "filing-sources@test.ru")
        applicant = client.post(
            "/api/v1/clients",
            json={"type": "individual", "full_name_or_company_name": "Иван Тестов"},
            headers=headers,
        ).json()
        application = client.post(
            "/api/v1/applications",
            json={"client_id": applicant["id"], "mark_name": "ТЕСТ", "mark_type": "word"},
            headers=headers,
        ).json()

        factory = async_sessionmaker(
            bind=async_engine, class_=AsyncSession, expire_on_commit=False
        )
        async with factory() as session:
            session.add(
                ExtractedField(
                    application_id=application["id"],
                    field_path="registry.sole_proprietor.full_name",
                    label="ФИО",
                    raw_value="Иван Тестов",
                    normalized_value="Иван Тестов",
                    extraction_method=ExtractionMethod.pdf_text_layer,
                    status=FieldStatus.matched,
                )
            )
            await session.commit()

        recorded = client.post(
            f"/api/v1/applications/{application['id']}/mark-details-suggestions",
            json={"description": "Словесное обозначение «ТЕСТ»."},
            headers=headers,
        )
        assert recorded.status_code == 204, recorded.text
        updated = client.put(
            f"/api/v1/applications/{application['id']}",
            json={"description_of_mark": "Словесное обозначение «ТЕСТ»."},
            headers=headers,
        )
        assert updated.status_code == 200, updated.text

        response = client.get(
            f"/api/v1/applications/{application['id']}/filing-package",
            headers=headers,
        )
        assert response.status_code == 200, response.text
        sources = {
            item["code"]: item for item in response.json()["field_sources"]["fields"]
        }
        assert sources["applicant_name"]["source"] == "document"
        assert sources["mark_description"]["source"] == "system"
        assert sources["mark_name"]["source"] == "user"
        assert sources["rospatent_application_number"]["source"] == "rospatent"

        client.put(
            f"/api/v1/applications/{application['id']}",
            json={"description_of_mark": "Исправлено пользователем"},
            headers=headers,
        )
        refreshed = client.get(
            f"/api/v1/applications/{application['id']}/filing-package",
            headers=headers,
        ).json()
        changed = {item["code"]: item for item in refreshed["field_sources"]["fields"]}
        assert changed["mark_description"]["source"] == "user"

    async def test_signatory_fields_with_date_can_be_saved(
        self, client, api_user_factory
    ):
        await api_user_factory("filing-signatory@test.ru", UserRole.client)
        headers = login_headers(client, "filing-signatory@test.ru")
        applicant = client.post(
            "/api/v1/clients",
            json={
                "type": "company",
                "full_name_or_company_name": 'ООО "Тест"',
            },
            headers=headers,
        ).json()
        application = client.post(
            "/api/v1/applications",
            json={"client_id": applicant["id"], "mark_name": "ТЕСТ", "mark_type": "word"},
            headers=headers,
        ).json()

        response = client.put(
            f"/api/v1/applications/{application['id']}",
            json={
                "signatory_name": "Алексеенко Андрей Сергеевич",
                "signatory_position": "Директор",
                "signature_date": "2026-08-25",
            },
            headers=headers,
        )

        assert response.status_code == 200, response.text
        assert response.json()["signatory_name"] == "Алексеенко Андрей Сергеевич"
        assert response.json()["signatory_position"] == "Директор"
        assert response.json()["signature_date"] == "2026-08-25"

    async def test_sound_mark_requires_audio_attachment(self, client, api_user_factory):
        await api_user_factory("filing-sound@test.ru", UserRole.client)
        headers = login_headers(client, "filing-sound@test.ru")
        applicant = client.post(
            "/api/v1/clients",
            json={"type": "individual", "full_name_or_company_name": "Иван Тестов"},
            headers=headers,
        ).json()
        application = client.post(
            "/api/v1/applications",
            json={"client_id": applicant["id"], "mark_name": "ДЖИНГЛ", "mark_type": "sound"},
            headers=headers,
        ).json()

        status = client.get(
            f"/api/v1/applications/{application['id']}/filing-package",
            headers=headers,
        )
        assert status.status_code == 200, status.text
        assert "mark_audio" in {item["code"] for item in status.json()["blockers"]}

    async def test_user_confirms_document_kind_before_packaging(self, client, api_user_factory):
        await api_user_factory("filing-kind@test.ru", UserRole.client)
        headers = login_headers(client, "filing-kind@test.ru")
        applicant = client.post(
            "/api/v1/clients",
            json={"type": "individual", "full_name_or_company_name": "Иван Тестов"},
            headers=headers,
        ).json()
        application = client.post(
            "/api/v1/applications",
            json={"client_id": applicant["id"], "mark_name": "ТЕСТ", "mark_type": "word"},
            headers=headers,
        ).json()
        upload = client.post(
            f"/api/v1/applications/{application['id']}/source-documents",
            files={
                "file": (
                    "doverennost.txt",
                    "ДОВЕРЕННОСТЬ\nНастоящей доверенностью заявитель уполномочивает представителя.",
                    "text/plain",
                )
            },
            headers=headers,
        )
        assert upload.status_code == 201, upload.text
        assert upload.json()["kind_requires_confirmation"] is True

        confirmed = client.put(
            f"/api/v1/source-documents/{upload.json()['id']}/kind",
            json={"document_kind": "power_of_attorney"},
            headers=headers,
        )
        assert confirmed.status_code == 200, confirmed.text
        assert confirmed.json()["document_kind"] == "power_of_attorney"
        assert confirmed.json()["kind_requires_confirmation"] is False

    async def test_passport_is_explicitly_excluded_from_filing_zip(self, client, api_user_factory):
        await api_user_factory("filing-passport@test.ru", UserRole.client)
        headers = login_headers(client, "filing-passport@test.ru")
        applicant = client.post(
            "/api/v1/clients",
            json={"type": "individual", "full_name_or_company_name": "Иван Тестов"},
            headers=headers,
        ).json()
        application = client.post(
            "/api/v1/applications",
            json={"client_id": applicant["id"], "mark_name": "ТЕСТ", "mark_type": "word"},
            headers=headers,
        ).json()
        upload = client.post(
            f"/api/v1/applications/{application['id']}/source-documents",
            files={
                "file": (
                    "passport.txt",
                    "ПАСПОРТ ГРАЖДАНИНА РОССИЙСКОЙ ФЕДЕРАЦИИ\nПаспорт выдан УМВД России",
                    "text/plain",
                )
            },
            headers=headers,
        )
        assert upload.status_code == 201, upload.text
        confirmed = client.put(
            f"/api/v1/source-documents/{upload.json()['id']}/kind",
            json={"document_kind": "passport"},
            headers=headers,
        )
        assert confirmed.status_code == 200, confirmed.text

        status = client.get(
            f"/api/v1/applications/{application['id']}/filing-package",
            headers=headers,
        ).json()

        assert status["excluded_documents"] == [
            {
                "filename": "passport.txt",
                "title": "Паспорт заявителя",
                "reason": (
                    "Хранится только в защищённом деле для сверки данных. "
                    "Копия паспорта не включается в ZIP и не направляется в Роспатент."
                ),
            }
        ]
        assert "passport.txt" not in {item["filename"] for item in status["documents"]}

    async def test_incomplete_case_returns_actionable_blockers(self, client, api_user_factory):
        await api_user_factory("filing-blocked@test.ru", UserRole.client)
        headers = login_headers(client, "filing-blocked@test.ru")
        applicant = client.post(
            "/api/v1/clients",
            json={"type": "individual", "full_name_or_company_name": "Иван Тестов"},
            headers=headers,
        ).json()
        application = client.post(
            "/api/v1/applications",
            json={"client_id": applicant["id"], "mark_name": "ТЕСТ", "mark_type": "word"},
            headers=headers,
        ).json()

        status = client.get(
            f"/api/v1/applications/{application['id']}/filing-package",
            headers=headers,
        )
        assert status.status_code == 200
        body = status.json()
        assert body["ready"] is False
        assert body["blockers"]
        assert body["requirements"]["version"] == "1.1.0"
        assert body["requirements"]["applicant_type"] == "individual"
        assert all({"title", "action", "section"} <= set(item) for item in body["blockers"])

        download = client.get(
            f"/api/v1/applications/{application['id']}/filing-package/download",
            headers=headers,
        )
        assert download.status_code == 409
        assert download.json()["detail"]["blockers"]

    async def test_ready_case_downloads_split_zip(self, client, api_user_factory, async_engine):
        user = await api_user_factory("filing-ready@test.ru", UserRole.client)
        factory = async_sessionmaker(
            bind=async_engine, class_=AsyncSession, expire_on_commit=False
        )
        async with factory() as session:
            applicant = Client(
                type=ClientType.company,
                full_name_or_company_name='ООО "Готовый пакет"',
                inn="7701234567",
                ogrn_or_ogrnip="1027700123456",
                address="123456, г. Москва, ул. Тестовая, д. 1",
                country="RU",
                email="owner@example.test",
                phone="+7 900 000-00-00",
                created_by_user_id=user.id,
            )
            session.add(applicant)
            await session.flush()
            application = TrademarkApplicationDraft(
                client_id=applicant.id,
                created_by_user_id=user.id,
                status=ApplicationStatus.memo_approved,
                mark_type=MarkType.word,
                mark_name="ГОТОВЫЙ ЗНАК",
                mark_text="ГОТОВЫЙ ЗНАК",
                description_of_mark="Словесное обозначение кириллицей",
                business_description="Разработка программного обеспечения",
                goods_services_raw="Программное обеспечение; разработка программ",
                territory="Российская Федерация",
                filing_method="electronic",
                signatory_name="Иванов Иван Иванович",
                signatory_position="генеральный директор",
                signature_date=date(2026, 8, 26),
            )
            session.add(application)
            await session.flush()
            session.add(
                NiceClassSuggestion(
                    application_id=application.id,
                    class_number=42,
                    class_description="разработка программного обеспечения",
                    rationale="Основной вид деятельности",
                    confidence=0.95,
                    category=NiceCategory.primary,
                    approved=True,
                    approved_by=user.id,
                )
            )
            for kind in (AnalysisKind.absolute_grounds, AnalysisKind.relative_grounds):
                session.add(
                    RiskAssessment(
                        application_id=application.id,
                        analysis_kind=kind,
                        overall_risk=RiskLevel.low,
                        summary="Существенные препятствия в доступном объёме проверки не выявлены.",
                        is_inconclusive=False,
                        search_mode=SearchMode.limited,
                        classes_considered_json=[42],
                        classes_confirmed=True,
                        requires_specialist_review=True,
                    )
                )
            session.add(
                AuditLog(
                    user_id=user.id,
                    application_id=application.id,
                    action="application.data.confirmed",
                    entity_type="TrademarkApplicationDraft",
                    entity_id=str(application.id),
                    new_value_json={"confirmed": True},
                )
            )
            await session.commit()
            application_id = application.id

        headers = login_headers(client, "filing-ready@test.ru")
        status = client.get(
            f"/api/v1/applications/{application_id}/filing-package",
            headers=headers,
        )
        assert status.status_code == 200, status.text
        body = status.json()
        assert body["ready"] is True, body["blockers"]
        assert body["filing_document_count"] == 1
        assert body["reference_document_count"] == 4
        assert body["class_numbers"] == [42]
        assert body["filing_fee"] > 0

        download = client.get(
            f"/api/v1/applications/{application_id}/filing-package/download",
            headers=headers,
        )
        assert download.status_code == 200, download.text
        assert download.headers["content-type"].startswith("application/zip")

        with zipfile.ZipFile(io.BytesIO(download.content)) as archive:
            names = set(archive.namelist())
            assert "01_ДЛЯ_ПОДАЧИ/01_заявление.docx" in names
            assert "01_ДЛЯ_ПОДАЧИ/02_перечень_товаров_и_услуг.docx" not in names
            assert "02_ДЛЯ_ВАС/01_инструкция_по_подаче.docx" in names
            assert "02_ДЛЯ_ВАС/02_расчёт_пошлин.docx" in names
            assert "02_ДЛЯ_ВАС/03_результат_проверки.docx" in names
            assert "02_ДЛЯ_ВАС/04_контрольный_список.txt" in names
            assert "README.txt" in names

            instruction = docx.Document(
                io.BytesIO(archive.read("02_ДЛЯ_ВАС/01_инструкция_по_подаче.docx"))
            )
            text = "\n".join(paragraph.text for paragraph in instruction.paragraphs)
            assert "не загружайте весь ZIP" in text
            assert "Сохраните подтверждение" in text
