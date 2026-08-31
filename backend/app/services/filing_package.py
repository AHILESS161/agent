"""Пакет для самостоятельной подачи заявки в Роспатент.

Архив намеренно разделён на две папки: документы, которые пользователь
переносит в официальный канал подачи, и справочные материалы для него самого.
Так расчёт пошлин или внутренний результат проверки невозможно случайно
приложить к заявлению как юридически значимый документ.
"""

from __future__ import annotations

import io
import re
import zipfile
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Any

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.infrastructure.database.models import (
    AnalysisKind,
    ClientRepresentative,
    DocumentKind,
    MarkType,
    RiskAssessment,
    SourceDocument,
    TrademarkApplicationDraft,
)
from app.services import file_storage
from app.services.application_draft import (
    DraftContent,
    collect_draft_content,
    load_mark_image_content,
    render_docx,
)
from app.services.class_analysis import load_class_context
from app.services.fee_calculator import calculate_trademark_fees
from app.services.filing_requirements import (
    filing_requirements_manifest,
    missing_required_items,
)
from app.services.field_provenance import field_sources_manifest

SERVICE_URL = (
    "https://rospatent.gov.ru/ru/stateservices/"
    "gosudarstvennaya-registraciya-tovarnogo-znaka-znaka-obsluzhivaniya-"
    "kollektivnogo-znaka-i-vydacha-svidetelstv-na-tovarnyy-znak-znak-"
    "obsluzhivaniya-kollektivnyy-znak-ih-dublikatov"
)
ARM_URL = "https://kpsrtz.fips.ru/"


class FilingPackageNotReady(ValueError):
    """Пакет нельзя выдавать как готовый, пока остались блокирующие пункты."""

    def __init__(self, blockers: list[dict[str, str]]) -> None:
        self.blockers = blockers
        super().__init__("Пакет документов пока не готов")


@dataclass
class FilingAttachment:
    filename: str
    title: str
    folder: str
    content: bytes


def _blocker(code: str, title: str, action: str, section: str) -> dict[str, str]:
    return {"code": code, "title": title, "action": action, "section": section}


def _safe_filename(name: str) -> str:
    clean = re.sub(r"[^\w.() -]+", "_", Path(name).name, flags=re.UNICODE).strip(" .")
    return clean[:180] or "document"


def _new_document(title: str, subtitle: str | None = None) -> docx.Document:
    document = docx.Document()
    for section in document.sections:
        section.top_margin = docx.shared.Cm(2)
        section.bottom_margin = docx.shared.Cm(2)
        section.left_margin = docx.shared.Cm(2.5)
        section.right_margin = docx.shared.Cm(2)
    heading = document.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if subtitle:
        paragraph = document.add_paragraph(subtitle)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return document


def _docx_bytes(document: docx.Document) -> bytes:
    payload = io.BytesIO()
    document.save(payload)
    return payload.getvalue()


def _add_label(document: docx.Document, label: str, value: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.add_run(f"{label}: ").bold = True
    paragraph.add_run(value or "—")


def _goods_document(application: TrademarkApplicationDraft, content: DraftContent) -> bytes:
    """Сформировать формальное приложение без справочного текста и декора."""
    document = docx.Document()
    for section in document.sections:
        section.top_margin = docx.shared.Cm(2)
        section.bottom_margin = docx.shared.Cm(2)
        section.left_margin = docx.shared.Cm(2)
        section.right_margin = docx.shared.Cm(2)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("ПЕРЕЧЕНЬ ТОВАРОВ И УСЛУГ")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)
    subtitle = document.add_paragraph(
        f"Приложение к заявке на обозначение "
        f"«{application.mark_text or application.mark_name or ''}»"
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Класс МКТУ"
    table.rows[0].cells[1].text = "Товары и (или) услуги"
    for number, description in content.classes:
        cells = table.add_row().cells
        cells[0].text = number
        cells[1].text = description or ""
    document.add_paragraph()
    signature = document.add_paragraph()
    if application.filing_method == "paper":
        signature.add_run("Подпись: __________________  ")
    signature.add_run(f"ФИО: {application.signatory_name or '__________________'}")
    if application.signatory_position:
        signature.add_run(f"  Должность: {application.signatory_position}")
    signed_at = (
        application.signature_date.strftime("%d.%m.%Y")
        if application.signature_date
        else "__________________"
    )
    document.add_paragraph(f"Дата подписания: {signed_at}")
    return _docx_bytes(document)


def _needs_goods_attachment(content: DraftContent) -> bool:
    """Нужен ли отдельный лист, если перечень не помещается в поле (511).

    В официальном бланке доступно десять строк. Оценка учитывает переносы
    длинных формулировок; короткий перечень из одного класса остаётся только в
    основном заявлении и не дублируется отдельным файлом.
    """
    used_lines = 0
    for _, description in content.classes:
        normalized = " ".join((description or "").split())
        used_lines += max(1, (len(normalized) + 109) // 110)
    return len(content.classes) > 10 or used_lines > 10


def _fees_document(fees: dict[str, Any]) -> bytes:
    document = _new_document(
        "РАСЧЁТ ПОШЛИН",
        f"Расчёт на {fees['calculated_at']} для {fees['class_count']} кл. МКТУ",
    )
    document.add_paragraph(
        "Это справочный расчёт для заявителя. Не прикладывайте его к заявлению вместо "
        "платёжного документа. Перед оплатой проверьте сумму в официальном сервисе."
    )
    for payment in fees.get("payments", []):
        document.add_heading(payment["title"], level=1)
        _add_label(
            document,
            "Основание тарифа",
            f"подпункт {payment['code']} приложения № 1 к Положению о патентных и иных пошлинах",
        )
        _add_label(document, "Когда оплачивать", payment["when"])
        _add_label(document, "Сумма", f"{payment['amount']:,} ₽".replace(",", " "))
    _add_label(
        document,
        "При подаче заявки",
        f"{fees['filing_total']:,} ₽".replace(",", " "),
    )
    _add_label(
        document,
        "Всего при положительном решении",
        f"{fees.get('total_selected') or fees['total_electronic']:,} ₽".replace(",", " "),
    )
    document.add_paragraph()
    document.add_paragraph(f"Официальная таблица пошлин: {fees['source_url']}")
    for warning in fees.get("warnings", []):
        document.add_paragraph(warning, style="List Bullet")
    return _docx_bytes(document)


_RISK_LABELS = {
    "low": "низкий",
    "medium": "средний",
    "high": "высокий",
    "critical": "критический",
}


def _client_analysis_summary(kind: str, assessment: RiskAssessment) -> str:
    """Свести юридический результат без внутренних ошибок интеграций."""
    if assessment.is_inconclusive:
        if kind == AnalysisKind.absolute_grounds.value:
            return (
                "Проверку самого обозначения пока не удалось завершить. "
                "Повторите её перед подачей; если результат снова не появится, "
                "передайте обозначение специалисту для ручной оценки."
            )
        return (
            "Поиск сходных обозначений пока не завершён. Повторите поиск перед "
            "подачей или передайте его специалисту."
        )
    return assessment.summary or (
        "По доступным данным обстоятельств, требующих отдельного предупреждения, не выявлено."
    )


def _analysis_document(application: TrademarkApplicationDraft, assessments: dict[str, RiskAssessment]) -> bytes:
    document = _new_document(
        "РЕЗУЛЬТАТ ПРЕДВАРИТЕЛЬНОЙ ПРОВЕРКИ",
        f"Обозначение «{application.mark_text or application.mark_name or ''}»",
    )
    document.add_paragraph(
        "Этот документ предназначен для заявителя и не подаётся в Роспатент. "
        "Проверка носит информационный характер и не гарантирует регистрацию."
    )
    labels = {
        AnalysisKind.absolute_grounds.value: "Проверка самого обозначения",
        AnalysisKind.relative_grounds.value: "Поиск сходных знаков и заявок",
    }
    for key, assessment in assessments.items():
        document.add_heading(labels.get(key, key), level=1)
        _add_label(
            document,
            "Уровень риска",
            _RISK_LABELS.get(assessment.overall_risk.value, assessment.overall_risk.value)
            if assessment.overall_risk
            else "не определён",
        )
        document.add_paragraph(_client_analysis_summary(key, assessment))
    return _docx_bytes(document)


def _instruction_document(
    application: TrademarkApplicationDraft,
    fees: dict[str, Any],
    filing_files: list[str],
    warnings: list[str],
) -> bytes:
    document = _new_document(
        "ИНСТРУКЦИЯ ПО ПОДАЧЕ ЗАЯВКИ",
        f"Дело № {application.id} · «{application.mark_text or application.mark_name or ''}»",
    )
    intro = document.add_paragraph()
    intro.add_run("Главное: ").bold = True
    intro.add_run(
        "не загружайте весь ZIP в Роспатент. Для подачи предназначены только файлы "
        "из папки «01_ДЛЯ_ПОДАЧИ». Папка «02_ДЛЯ_ВАС» содержит подсказки и расчёты."
    )

    steps = [
        (
            "1. Проверьте пакет",
            "Откройте заявление и перечень товаров и услуг. Сверьте ФИО или наименование, "
            "адрес, идентификаторы, обозначение и классы МКТУ. Не меняйте перечень после "
            "проверки без повторного расчёта пошлины и поиска сходных знаков.",
        ),
        (
            "2. Подготовьте официальный канал",
            f"Откройте АРМ «Регистратор»: {ARM_URL}. Войдите способом, который предлагает "
            "официальный сервис, и создайте заявку на товарный знак.",
        ),
        (
            "3. Перенесите сведения",
            "Перенесите данные из заявления в поля официальной формы. Загрузите изображение "
            "обозначения и доверенность только если они есть в папке для подачи. Отдельный "
            "перечень используйте, когда список не помещается в форме.",
        ),
        (
            "4. Оплатите первый этап",
            (
                f"К подаче рассчитано {fees['filing_total']:,} ₽. Оплачиваются регистрация заявки "
                "и экспертиза обозначения. Пошлину за регистрацию знака сейчас не платите — она "
                "понадобится после положительного решения."
            ).replace(",", " "),
        ),
        (
            "5. Выполните проверку и подписание",
            "Исправьте все ошибки форматно-логического контроля официального сервиса. "
            "Проверьте сформированный бланк, подпишите документы электронной подписью, если "
            "этого требует выбранный канал, и только после этого отправьте пакет.",
        ),
        (
            "6. Сохраните подтверждение",
            "Скачайте квитанцию о приёме, запишите номер заявки и дату подачи. Эти сведения "
            "определяют дальнейшую переписку и приоритет.",
        ),
        (
            "7. Следите за сообщениями",
            "Регулярно проверяйте корреспонденцию Роспатента. На запросы и уведомления нужно "
            "отвечать в указанный срок. Не оплачивайте второй этап до соответствующего решения.",
        ),
    ]
    for heading, text in steps:
        document.add_heading(heading, level=1)
        document.add_paragraph(text)

    document.add_heading("Файлы для подачи", level=1)
    for filename in filing_files:
        document.add_paragraph(filename, style="List Bullet")

    if warnings:
        document.add_heading("На что обратить внимание", level=1)
        for warning in warnings:
            document.add_paragraph(warning, style="List Bullet")

    document.add_heading("Официальные ссылки", level=1)
    document.add_paragraph(f"Описание государственной услуги: {SERVICE_URL}")
    document.add_paragraph(f"АРМ «Регистратор»: {ARM_URL}")
    document.add_paragraph(f"Таблица пошлин: {fees['source_url']}")
    return _docx_bytes(document)


def _checklist_text(application: TrademarkApplicationDraft, filing_files: list[str]) -> bytes:
    lines = [
        "КОНТРОЛЬНЫЙ СПИСОК ПЕРЕД ОТПРАВКОЙ",
        f"Заявка № {application.id}: {application.mark_text or application.mark_name or ''}",
        "",
        "[ ] Наименование/ФИО и адрес заявителя проверены",
        "[ ] ИНН, ОГРН/ОГРНИП или данные физлица проверены",
        "[ ] Обозначение и его вид указаны правильно",
        "[ ] Перечень товаров и услуг совпадает с подтверждёнными классами МКТУ",
        "[ ] Сумма первого этапа повторно проверена перед оплатой",
        "[ ] Ошибки официального форматно-логического контроля исправлены",
        "[ ] Заявление и необходимые приложения подписаны",
        "[ ] После отправки сохранены номер заявки, дата и квитанция",
        "[ ] Настроено напоминание проверять корреспонденцию Роспатента",
        "",
        "В ПАПКЕ ДЛЯ ПОДАЧИ:",
        *[f"- {name}" for name in filing_files],
    ]
    return "\n".join(lines).encode("utf-8-sig")


async def _latest_assessments(
    session: AsyncSession, application_id: int
) -> dict[str, RiskAssessment]:
    result: dict[str, RiskAssessment] = {}
    for kind in (AnalysisKind.absolute_grounds, AnalysisKind.relative_grounds):
        assessment = (
            await session.execute(
                select(RiskAssessment)
                .where(
                    RiskAssessment.application_id == application_id,
                    RiskAssessment.analysis_kind == kind,
                )
                .order_by(RiskAssessment.id.desc())
                .limit(1)
            )
        ).scalar_one_or_none()
        if assessment:
            result[kind.value] = assessment
    return result


async def _filing_attachments(
    session: AsyncSession, application: TrademarkApplicationDraft
) -> tuple[list[FilingAttachment], list[SourceDocument]]:
    documents = list(
        (
            await session.execute(
                select(SourceDocument)
                .where(SourceDocument.application_id == application.id)
                .order_by(SourceDocument.id)
            )
        )
        .scalars()
        .all()
    )
    attachments: list[FilingAttachment] = []
    selected: list[SourceDocument] = []
    for document in documents:
        if document.kind_requires_confirmation:
            continue
        if document.document_kind is DocumentKind.mark_image and (
            application.mark_type not in {MarkType.figurative, MarkType.combined}
            or application.mark_image_file_id != str(document.id)
        ):
            # В пакет попадает только текущая версия изображения и только для
            # вида знака, которому графическое приложение действительно нужно.
            continue
        allowed = {
            DocumentKind.mark_image,
            DocumentKind.mark_audio,
            DocumentKind.power_of_attorney,
        }
        if application.priority_claim:
            # Для приоритета отдельного enum пока нет. Такой файл сохраняется
            # как «иной документ» и включается только когда приоритет заявлен.
            allowed.add(DocumentKind.other)
        if document.document_kind not in allowed:
            continue
        try:
            payload = file_storage.read_file(document.stored_path)
        except FileNotFoundError:
            continue
        prefix = (
            "03"
            if document.document_kind is DocumentKind.mark_image
            else "04"
            if document.document_kind is DocumentKind.mark_audio
            else "05"
            if document.document_kind is DocumentKind.power_of_attorney
            else "06"
        )
        name = f"{prefix}_{_safe_filename(document.original_filename)}"
        title = {
            DocumentKind.mark_image: "Изображение обозначения",
            DocumentKind.mark_audio: "Аудиозапись звукового обозначения",
            DocumentKind.power_of_attorney: "Доверенность представителя",
            DocumentKind.other: "Документ, подтверждающий приоритет",
        }[document.document_kind]
        attachments.append(FilingAttachment(name, title, "01_ДЛЯ_ПОДАЧИ", payload))
        selected.append(document)
    return attachments, documents


async def filing_package_status(
    session: AsyncSession, application: TrademarkApplicationDraft
) -> dict[str, Any]:
    content = await collect_draft_content(session, application)
    classes = await load_class_context(session, application.id)
    fees = await calculate_trademark_fees(session, application.id)
    assessments = await _latest_assessments(session, application.id)
    attachments, source_documents = await _filing_attachments(session, application)
    excluded_documents = [
        {
            "filename": document.original_filename,
            "title": "Паспорт заявителя",
            "reason": (
                "Хранится только в защищённом деле для сверки данных. "
                "Копия паспорта не включается в ZIP и не направляется в Роспатент."
            ),
        }
        for document in source_documents
        if document.document_kind is DocumentKind.passport
    ]
    representative = (
        (
            await session.execute(
                select(ClientRepresentative).where(
                    ClientRepresentative.id == application.representative_id,
                    ClientRepresentative.client_id == application.client_id,
                )
            )
        ).scalar_one_or_none()
        if application.representative_id
        else None
    )
    attachment_codes = {
        {
            "Изображение обозначения": "mark_image",
            "Аудиозапись звукового обозначения": "mark_audio",
            "Доверенность представителя": "power_of_attorney",
            "Документ, подтверждающий приоритет": "priority_proof",
        }[item.title]
        for item in attachments
        if item.title in {
            "Изображение обозначения",
            "Аудиозапись звукового обозначения",
            "Доверенность представителя",
            "Документ, подтверждающий приоритет",
        }
    }
    requirements = filing_requirements_manifest(
        application,
        has_representative=representative is not None,
        representative=representative,
        available_attachments=attachment_codes,
    )
    field_sources = await field_sources_manifest(session, application)

    blockers: list[dict[str, str]] = []
    from app.services.data_confirmation import data_confirmation_state

    confirmation = await data_confirmation_state(session, application)
    if not confirmation["confirmed"]:
        blockers.append(
            _blocker(
                "data_confirmation",
                "Подтверждение сведений",
                "Сверьте реквизиты и данные знака, затем подтвердите их",
                "check",
            )
        )
    for item in missing_required_items(requirements):
        blockers.append(
            _blocker(item["code"], item["title"], item["action"], item["section"])
        )
    if not classes.is_confirmed:
        blockers.append(
            _blocker("classes", "Классы МКТУ", "Подтвердите хотя бы один класс", "check")
        )
    expected_analysis = {kind.value for kind in (AnalysisKind.absolute_grounds, AnalysisKind.relative_grounds)}
    for missing in sorted(expected_analysis - assessments.keys()):
        label = "Проверка самого обозначения" if missing == AnalysisKind.absolute_grounds.value else "Поиск сходных знаков"
        blockers.append(_blocker("analysis", label, "Завершите проверку", "check"))
    analysis_warnings: list[str] = []
    for assessment in assessments.values():
        # Оба анализа должны быть запущены, но неопределённый результат одной
        # предварительной проверки не должен навсегда блокировать документы.
        # Он переносится в пакет как предупреждение. Текущая готовность классов
        # проверяется выше, а не по снимку внутри старого assessment.
        if assessment.is_inconclusive:
            analysis_warnings.append(
                assessment.inconclusive_reason
                or "Одна из предварительных проверок не дала надёжного вывода"
            )
    if not fees.get("can_calculate"):
        blockers.append(_blocker("fees", "Расчёт пошлин", "Подтвердите классы МКТУ", "fees"))

    # Одинаковая причина из формы и явных правил показывается один раз.
    unique: dict[tuple[str, str], dict[str, str]] = {}
    for item in blockers:
        unique[(item["code"], item["title"])] = item
    blockers = list(unique.values())

    warnings = [
        "Пакет подготовлен для обычной самостоятельной подачи одного товарного знака.",
        "Перед отправкой ещё раз проверьте данные в официальном сервисе Роспатента.",
    ]
    warnings.extend(analysis_warnings)
    risk_order = {"low": 0, "medium": 1, "high": 2, "critical": 3}
    risk_levels = [a.overall_risk.value for a in assessments.values() if a.overall_risk]
    overall_risk = max(risk_levels, key=risk_order.get) if risk_levels else None
    if overall_risk in {"high", "critical"}:
        warnings.insert(
            0,
            "Проверка выявила высокий риск отказа. Пакет можно скачать, но до подачи рекомендуется изменить знак или обратиться к специалисту.",
        )
    if application.priority_claim:
        warnings.append("Заявлен приоритет: проверьте вид, дату, номер и срок представления подтверждающего документа.")

    include_goods_attachment = _needs_goods_attachment(content)
    manifest = [
        {
            "filename": "01_заявление.docx",
            "title": "Заявление на регистрацию товарного знака",
            "folder": "01_ДЛЯ_ПОДАЧИ",
            "purpose": "Перенесите сведения в официальный сервис или используйте сформированный бланк",
        },
    ]
    if include_goods_attachment:
        manifest.append(
            {
                "filename": "02_перечень_товаров_и_услуг.docx",
                "title": "Перечень товаров и услуг по классам МКТУ",
                "folder": "01_ДЛЯ_ПОДАЧИ",
                "purpose": "Продолжение перечня из поля (511) заявления",
            }
        )
    manifest.extend(
        {
            "filename": item.filename,
            "title": item.title,
            "folder": item.folder,
            "purpose": "Условное приложение к заявке",
        }
        for item in attachments
    )
    manifest.extend(
        [
            {
                "filename": "01_инструкция_по_подаче.docx",
                "title": "Пошаговая инструкция по подаче",
                "folder": "02_ДЛЯ_ВАС",
                "purpose": "Не загружать в Роспатент",
            },
            {
                "filename": "02_расчёт_пошлин.docx",
                "title": "Расчёт пошлин по выбранным классам",
                "folder": "02_ДЛЯ_ВАС",
                "purpose": "Не загружать в Роспатент",
            },
            {
                "filename": "03_результат_проверки.docx",
                "title": "Результат предварительной проверки",
                "folder": "02_ДЛЯ_ВАС",
                "purpose": "Не загружать в Роспатент",
            },
            {
                "filename": "04_контрольный_список.txt",
                "title": "Чек-лист перед отправкой",
                "folder": "02_ДЛЯ_ВАС",
                "purpose": "Отметьте шаги вручную",
            },
        ]
    )

    return {
        "application_id": application.id,
        "ready": not blockers,
        "blockers": blockers,
        "requirements": requirements,
        "field_sources": field_sources,
        "warnings": warnings,
        "documents": manifest,
        "excluded_documents": excluded_documents,
        "filing_document_count": sum(1 for item in manifest if item["folder"] == "01_ДЛЯ_ПОДАЧИ"),
        "reference_document_count": sum(1 for item in manifest if item["folder"] == "02_ДЛЯ_ВАС"),
        "class_numbers": [item.class_number for item in classes.approved],
        "overall_risk": overall_risk,
        "filing_fee": fees.get("filing_total"),
        "registration_fee": fees.get("registration_total"),
        "total_fee": fees.get("total_selected") or fees.get("total_electronic"),
        "generated_for": date.today().isoformat(),
        "_content": content,
        "_fees": fees,
        "_assessments": assessments,
        "_attachments": attachments,
        "_include_goods_attachment": include_goods_attachment,
    }


def public_filing_package_status(status: dict[str, Any]) -> dict[str, Any]:
    """Убрать служебные ORM-объекты перед возвратом JSON клиенту."""
    return {key: value for key, value in status.items() if not key.startswith("_")}


async def render_filing_package(
    session: AsyncSession, application: TrademarkApplicationDraft
) -> tuple[bytes, dict[str, Any]]:
    status = await filing_package_status(session, application)
    if not status["ready"]:
        raise FilingPackageNotReady(status["blockers"])

    content: DraftContent = status.pop("_content")
    fees: dict[str, Any] = status.pop("_fees")
    assessments: dict[str, RiskAssessment] = status.pop("_assessments")
    attachments: list[FilingAttachment] = status.pop("_attachments")
    include_goods_attachment: bool = status.pop("_include_goods_attachment")
    filing_names = [item["filename"] for item in status["documents"] if item["folder"] == "01_ДЛЯ_ПОДАЧИ"]

    output = io.BytesIO()
    with zipfile.ZipFile(output, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "01_ДЛЯ_ПОДАЧИ/01_заявление.docx",
            render_docx(
                content,
                application,
                mark_image=await load_mark_image_content(session, application),
                include_goods_attachment=include_goods_attachment,
            ),
        )
        if include_goods_attachment:
            archive.writestr(
                "01_ДЛЯ_ПОДАЧИ/02_перечень_товаров_и_услуг.docx",
                _goods_document(application, content),
            )
        for attachment in attachments:
            archive.writestr(f"{attachment.folder}/{attachment.filename}", attachment.content)
        archive.writestr(
            "02_ДЛЯ_ВАС/01_инструкция_по_подаче.docx",
            _instruction_document(application, fees, filing_names, status["warnings"]),
        )
        archive.writestr("02_ДЛЯ_ВАС/02_расчёт_пошлин.docx", _fees_document(fees))
        archive.writestr(
            "02_ДЛЯ_ВАС/03_результат_проверки.docx",
            _analysis_document(application, assessments),
        )
        archive.writestr(
            "02_ДЛЯ_ВАС/04_контрольный_список.txt",
            _checklist_text(application, filing_names),
        )
        archive.writestr(
            "README.txt",
            (
                "ПАКЕТ ДЛЯ САМОСТОЯТЕЛЬНОЙ ПОДАЧИ\n\n"
                "В Роспатент загружаются только применимые файлы из папки "
                "01_ДЛЯ_ПОДАЧИ. Папка 02_ДЛЯ_ВАС содержит инструкцию, расчёт и "
                "результат проверки; её содержимое не является приложением к заявке.\n\n"
                "Перед отправкой откройте 02_ДЛЯ_ВАС/01_инструкция_по_подаче.docx."
            ).encode("utf-8-sig"),
        )
    return output.getvalue(), status
